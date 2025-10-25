from __future__ import annotations

import io
import math
import re
import numpy as np
import pandas as pd
import streamlit as st
import matplotlib.pyplot as plt
from matplotlib.gridspec import GridSpec
from matplotlib.ticker import FuncFormatter, MaxNLocator
from mplsoccer import Radar
from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor, Mm, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from datetime import datetime, date


OUTPUT_FILE = "nassim_ait_mouhou_executive_report.docx"
REPORT_TITLE = "Performance radar and percentile overview"
REPORT_SUBTITLE = "Executive report"
PLAYER_NAME_DEFAULT = "Nassim Ait Mouhou"
REPORT_AUTHOR = "Joao Alberto Kolling"
REPORT_CONFIDENTIALITY_NOTE = "Confidential — For internal club use only."
PLAYER_CLUB_DEFAULT = "VVV Venlo"
PLAYER_POSITION_DEFAULT = "LAMF, LW"
PLAYER_AGE_DEFAULT = 21
SAMPLE_SIZE_DEFAULT = 486
COMPETITION_LEVEL_DEFAULT = "—"
REPORT_ID = "VIF-PR-2025-10-24"
REPORT_VERSION = "v1.0"
A4_WIDTH_MM = 210
A4_HEIGHT_MM = 297
PAGE_MARGIN_MM = 20
HEADER_FOOTER_DISTANCE_MM = 12

METHODOLOGY_COHORT = 435
METHODOLOGY_WINDOW = "Rolling 18-month window (Aug 2023 – Feb 2025)"
METHODOLOGY_REVERSALS = "Negative-impact metrics are inverted before percentile ranking."

STANDOUTS_DEFAULT = [
    {
        "metric": "Progressive runs per 90",
        "descriptor": "98th pct (5.13)",
        "pct": 98.0,
        "value": 5.13,
        "median": None,
    },
    {
        "metric": "Progression",
        "descriptor": "95th pct (48.6)",
        "pct": 95.0,
        "value": 48.6,
        "median": None,
    },
    {
        "metric": "Offensive Intensity",
        "descriptor": "94th pct (6.22)",
        "pct": 94.0,
        "value": 6.22,
        "median": None,
    },
]

DEVELOPMENT_AREAS_DEFAULT = [
    {
        "metric": "Passing Quality",
        "descriptor": "29th pct (49.1)",
        "pct": 29.0,
        "value": 49.1,
        "median": None,
    },
]

def _player_age(row):
    # Try common age fields
    for key in ["Age", "age"]:
        if key in row and row[key] is not None:
            try:
                val = int(float(row[key]))
                if val > 0:
                    return val
            except Exception:
                pass
    # Try DOB parsing (YYYY-MM-DD or DD/MM/YYYY etc.)
    for key in ["DOB", "Date of Birth", "Birthdate", "Birth Date"]:
        if key in row and isinstance(row[key], str) and row[key].strip():
            txt = row[key].strip()
            # try multiple formats
            fmts = ["%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%m/%d/%Y", "%Y/%m/%d"]
            for fmt in fmts:
                try:
                    dob = datetime.strptime(txt, fmt).date()
                    today = date.today()
                    age = today.year - dob.year - ((today.month, today.day) < (dob.month, dob.day))
                    if age > 0:
                        return age
                except Exception:
                    continue
    # Try Birth Year
    for key in ["Birth Year", "Birth year", "YOB", "Year of Birth"]:
        if key in row and row[key] is not None:
            try:
                year = int(float(row[key]))
                today = date.today()
                age = today.year - year
                if 10 < age < 60:
                    return age
            except Exception:
                pass
    return None



def _fix_npxg_block(df):
    import numpy as np
    if "xG" in df.columns and "Penalties taken" in df.columns:
        df["npxG"] = df["xG"].fillna(0) - df["Penalties taken"].fillna(0) * 0.81
    # per 90
    if "npxG" in df.columns and "Minutes played" in df.columns:
        with np.errstate(divide="ignore", invalid="ignore"):
            mp = df["Minutes played"].replace(0, np.nan).astype(float)
            df["npxG per 90"] = (df["npxG"].astype(float) / mp) * 90.0
    # per shot
    if "npxG" in df.columns and "Shots" in df.columns:
        with np.errstate(divide="ignore", invalid="ignore"):
            df["npxG per Shot"] = df["npxG"].astype(float) / df["Shots"].replace(0, np.nan).astype(float)
    # G-xG
    if "Goals" in df.columns and "xG" in df.columns:
        df["G-xG"] = df["Goals"].fillna(0) - df["xG"].fillna(0)
    return df


# ===================== CONFIG & STYLE =====================
st.set_page_config(
    page_title="Composite Metrics & Radar",
    page_icon="⚽",
    layout="wide",
    initial_sidebar_state="expanded",
)
st.markdown(
    """
    <style>
    .metric-card {padding: 1rem; border: 1px solid rgba(49,51,63,0.2); border-radius: 12px;}
    .section {margin-top: .75rem; margin-bottom: .25rem; font-weight: 600; opacity: .9}
    .subtle {color: rgba(49,51,63,0.6)}
    .stTabs [data-baseweb="tab-list"] {gap: 12px}
    .stTabs [data-baseweb="tab"] {background: #f6f6f9; padding: 8px 14px; border-radius: 10px;}
    .small {font-size: 0.9rem}
    </style>
    """,
    unsafe_allow_html=True,
)

# ===================== CONSTANTS =====================
OFFENSIVE_COMPONENTS = [
    "Successful attacking actions per 90",
    "xG per 90",
    "xA per 90",
    "Key passes per 90",
    "Deep completions per 90",
    "Deep completed crosses per 90",
    "Progressive runs per 90",
    "Passes to penalty area per 90",
    "Smart passes per 90",
    "Crosses to goalie box per 90",
    "Touches in box per 90",
]
DEFAULT_OFF_WEIGHTS = {
    "Successful attacking actions per 90": 1.0,
    "xG per 90": 1.0,
    "xA per 90": 1.0,
    "Key passes per 90": 1.0,
    "Deep completions per 90": 1.0,
    "Deep completed crosses per 90": 0.8,
    "Progressive runs per 90": 0.6,
    "Passes to penalty area per 90": 0.8,
    "Smart passes per 90": 0.5,
    "Crosses to goalie box per 90": 0.4,
    "Touches in box per 90": 0.4,
}
PHYS_COLS = {
    "distance_total_m": "Distance P90",
    "running_m": "Running Distance P90",
    "hi_m": "HI Distance P90",
    "hsr_m": "HSR Distance P90",
    "sprint_m": "Sprint Distance P90",
    "hsr_cnt": "HSR Count P90",
    "sprint_cnt": "Sprint Count P90",
    "expl_hsr_cnt": "Explosive Acceleration to HSR Count P90",
    "expl_sprint_cnt": "Explosive Acceleration to Sprint Count P90",
}
DEF_COLS = [
    "Successful defensive actions per 90",
    "Defensive duels per 90",
    "Defensive duels won, %",
    "Aerial duels per 90",
    "Aerial duels won, %",
    "PAdj Sliding tackles",
    "PAdj Interceptions",
    "Shots blocked per 90",
]
GK_METRICS = [
    "Save rate, %","Prevented goals per 90","Conceded goals per 90","Shots against per 90","Clean sheets",
    "Back passes received as GK per 90","xG against","xG against per 90","Prevented goals","Shots against",
    "Conceded goals","Exits per 90",
]
ID_COLS_CANDIDATES = ["Player", "Short Name", "Team", "Position", "Minutes played", "Minutes"]

PRESETS = {
    "forward": [
        "npxG per 90","xG per 90","Shots per 90","Shots on target, %",
        "xA per 90","Key passes per 90","Touches in box per 90",
        "Progressive runs per 90","Deep completions per 90",
        "Finishing","Poaching","Aerial Threat",
        "Work Rate Offensive","Offensive Intensity","Offensive Explosion",
    ],
    "winger": [
        "Dribbles per 90","Dribbles won, %","Crosses per 90","Accurate crosses, %",
        "Deep completed crosses per 90","Progressive runs per 90","xA per 90","Key passes per 90",
        "Creativity","Progression",
        "Work Rate Offensive","Offensive Intensity","Offensive Explosion",
    ],
    "attacking_midfielder": [
        "Creativity","Progression","xA per 90","Key passes per 90","Smart passes per 90",
        "Deep completions per 90","xG Buildup",
        "Work Rate Offensive","Offensive Intensity","Offensive Explosion",
    ],
    "central_midfielder": [
        "Progression","Passing Quality","Creativity","xG Buildup",
        "PAdj Interceptions","Successful defensive actions per 90",
        "Work Rate Defensive","Defensive Intensity",
    ],
    "defensive_midfielder": [
        "Defence","Progression","Passing Quality","Discipline","Involvement",
        "xG Buildup","PAdj Interceptions","Successful defensive actions per 90",
        "Work Rate Defensive","Defensive Intensity","Defensive Explosion",
    ],
    "full_back": [
        "Progression","Creativity","Passing Quality","Defence","Aerial Defence",
        "Deep completed crosses per 90","Progressive runs per 90","Crosses per 90",
        "Work Rate Defensive","Defensive Intensity","Defensive Explosion",
    ],
    "center_back": [
        "Defence","Aerial Defence","Aerial duels per 90","Aerial duels won, %",
        "Defensive duels per 90","Defensive duels won, %","PAdj Interceptions","PAdj Sliding tackles",
        "Passes to final third per 90","Accurate passes %",
        "Work Rate Defensive","Defensive Intensity","Defensive Explosion",
    ],
    "goalkeeper": GK_METRICS + [
        "Passes per 90","Accurate passes, %","Long passes per 90","Accurate long passes, %","Aerial duels per 90","Aerial duels won, %"
    ],
    "general_summary": ["Involvement","Box Threat","Creativity","xG Buildup","Progression","Defence","Discipline","Passing Quality"],
    "defensive_actions": ["Defence","Aerial Defence","PAdj Interceptions","Successful defensive actions per 90","Defensive duels won, %","Shots blocked per 90","Discipline"],
    "playmaking": ["Creativity","Passing Quality","Progression","xG Buildup","Successful attacking actions per 90","Deep completions per 90","Involvement","Key passes per 90"],
    "aerial_duels": ["Aerial Threat","Aerial Defence","Aerial duels per 90","Aerial duels won, %","Head goals per 90","Involvement","Defence"],
    "shooting": ["npxG per 90","npxG per Shot","Finishing","Goal conversion, %","Shots on target, %","G-xG","Box Threat"],
    "counter_attack": ["Progression","Accelerations per 90","Dribbles per 90","Successful dribbles, %","npxG per 90","Finishing","Box Threat"],
    "build_up": ["Passing Quality","Progression","xG Buildup","Deep completions per 90","Involvement","Creativity","Discipline"],
    "crossing": ["Crossing","Accurate crosses, %","Deep completed crosses per 90","xA per 90","Shot assists per 90","Creativity","Passing Quality"],
    "striker": ["npxG per 90","npxG per Shot","Finishing","Poaching","Aerial Threat","Box Threat","Involvement","Touches in box per 90"],
}

NEGATE_METRICS = {
    "Conceded goals per 90": True,
    "xG against per 90": True,
    "Fouls per 90": True,
    "Yellow cards per 90": True,
    "Red cards per 90": True,
}

# ===================== HELPERS =====================
@st.cache_data(show_spinner=False)
def _load_excel(file: io.BytesIO) -> pd.DataFrame:
    try:
        return pd.read_excel(file, sheet_name=0)
    except Exception:
        file.seek(0)
        return pd.read_excel(file)

def _ensure_cols(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    for c in cols:
        if c not in df.columns:
            df[c] = np.nan
    return df


def _slugify_filename(text: str) -> str:
    slug = re.sub(r"[^0-9A-Za-z_-]+", "_", (text or "")).strip("_")
    slug = slug.lower()
    return slug[:60] or "relatorio"

def _safe_s(df: pd.DataFrame, col: str) -> pd.Series:
    return df[col].astype(float).fillna(0) if col in df.columns else pd.Series(0.0, index=df.index)

def _zscore(series: pd.Series) -> pd.Series:
    s = pd.to_numeric(series, errors="coerce")
    m = s.mean(skipna=True)
    sd = s.std(skipna=True, ddof=0)
    if not np.isfinite(sd) or sd == 0:
        return pd.Series(0.0, index=s.index)
    return (s - m) / sd

# ===================== DERIVED METRICS =====================
def compute_derived_metrics(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    needed = set([
        "xG","Penalties taken","Goals","Shots","Touches in box per 90",
        "xA per 90","Shot assists per 90","Key passes per 90","Deep completions per 90",
        "Deep completed crosses per 90","Accurate passes %","Accurate passes, %",
        "Smart passes per 90","Through passes per 90","Passes to penalty area per 90",
        "Accurate smart passes, %","Accurate through passes, %","Accurate passes to penalty area, %",
        "Progressive passes per 90","Progressive runs per 90","Dribbles per 90","Accelerations per 90",
        "Accurate progressive passes, %","Successful dribbles, %",
        "Successful defensive actions per 90","PAdj Interceptions","Shots blocked per 90","PAdj Sliding tackles",
        "Defensive duels per 90","Defensive duels won, %","Aerial duels won, %","Aerial duels per 90",
        "Passes to final third per 90","Accurate passes to final third, %","Forward passes per 90","Accurate forward passes, %",
        "Long passes per 90","Accurate long passes, %","Lateral passes per 90","Accurate lateral passes, %",
        "Back passes per 90","Accurate back passes, %","Passes per 90",
        "Head goals per 90","Goal conversion, %","Received passes per 90",
        "Fouls per 90","Yellow cards per 90","Red cards per 90",
        "Minutes played"
    ])
    _ensure_cols(df, list(needed))

    # npxG & friends
    if "xG" in df and "Penalties taken" in df:
        df["npxG"] = (df["xG"] - (df["Penalties taken"].fillna(0) * 0.81))
    if "Goals" in df and "xG" in df:
        df["G-xG"] = _zscore(df["Goals"] - df["xG"])
    if "npxG" in df and "Minutes played" in df:
        with np.errstate(divide="ignore", invalid="ignore"):
            df["npxG per 90"] = df["npxG"] / (df["Minutes played"].replace(0, np.nan) / 90)
    if "npxG" in df and "Shots" in df:
        with np.errstate(divide="ignore", invalid="ignore"):
            df["npxG per Shot"] = df["npxG"] / df["Shots"].replace(0, np.nan)
    # Box Threat
    if "npxG per 90" in df and "Touches in box per 90" in df:
        with np.errstate(divide="ignore", invalid="ignore"):
            denom = np.log(df["Touches in box per 90"].fillna(0) + 1)
            denom = denom.replace(0, np.nan)
            df["Box Threat"] = _zscore(df["npxG per 90"] / denom)

    # xG Buildup
    xgb_w = {
        "xA per 90": 3.0, "Shot assists per 90": 3.0, "npxG per 90": 2.5, "Key passes per 90": 2.5,
        "Deep completions per 90": 2.5, "Deep completed crosses per 90": 2.0, "Second assists per 90": 1.5,
        "Accurate passes %": 1.0,
    }
    parts = []
    total_w = 0.0
    for col, w in xgb_w.items():
        if col in df.columns:
            s = df[col].fillna(0)
            if "%" in col and s.max() > 1:
                s = s / 100.0
            parts.append(_zscore(s) * w)
            total_w += w
    if parts and total_w > 0:
        df["xG Buildup"] = sum(parts) / total_w

    # Creativity
    vol = {"Smart passes per 90": 1.0, "Through passes per 90": 0.8, "Passes to penalty area per 90": 0.6}
    acc = {"Accurate smart passes, %": 1.0, "Accurate through passes, %": 0.9, "Accurate passes to penalty area, %": 0.8, "Accurate passes %": 0.6}
    vol_scores, acc_scores = [], []
    for c, w in vol.items():
        if c in df.columns: vol_scores.append(_zscore(df[c].fillna(0)) * w)
    for c, w in acc.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if s.max() > 1: s = s/100.0
            acc_scores.append(_zscore(s) * w)
    if vol_scores and acc_scores:
        df["Creativity"] = 0.6 * np.nanmean(vol_scores, axis=0) + 0.4 * np.nanmean(acc_scores, axis=0)

    # Progression
    prog_vol = {"Progressive passes per 90": 1.0, "Progressive runs per 90": 0.9, "Dribbles per 90": 0.7, "Accelerations per 90": 0.6}
    prog_acc = {"Accurate progressive passes, %": 1.0, "Successful dribbles, %": 0.8, "Accurate passes %": 0.3}
    vol_scores, acc_scores = [], []
    for c, w in prog_vol.items():
        if c in df.columns: vol_scores.append(_zscore(df[c].fillna(0)) * w)
    for c, w in prog_acc.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if s.max() > 1: s = s/100.0
            acc_scores.append(_zscore(s) * w)
    if vol_scores and acc_scores:
        df["Progression"] = 0.6 * np.nanmean(vol_scores, axis=0) + 0.4 * np.nanmean(acc_scores, axis=0)

    # Defence
    def_w = {
        "Successful defensive actions per 90": 1.5, "PAdj Interceptions": 1.5, "Shots blocked per 90": 1.0,
        "PAdj Sliding tackles": 1.0, "Defensive duels per 90": 0.5, "Defensive duels won, %": 3.0, "Aerial duels won, %": 1.5,
    }
    parts = []; total_w = 0.0
    for c, w in def_w.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if "%" in c and s.max() > 1: s = s/100.0
            parts.append(_zscore(s) * w); total_w += w
    if parts and total_w > 0:
        df["Defence"] = sum(parts) / total_w

    # Involvement
    inv_w = {
        "Passes per 90": .20, "Received passes per 90": .15, "Touches per 90": .05,
        "Defensive duels per 90": .10, "PAdj Interceptions": .10, "Successful defensive actions per 90": .05,
        "Touches in box per 90": .10, "Offensive duels per 90": .10, "Progressive runs per 90": .05,
        "Aerial duels per 90": .10,
    }
    parts = []
    for c, w in inv_w.items():
        if c in df.columns: parts.append(_zscore(df[c].fillna(0)) * w)
    if parts: df["Involvement"] = sum(parts)

    # Discipline (negative)
    if all(c in df.columns for c in ["Fouls per 90","Yellow cards per 90","Red cards per 90"]):
        penalty = (df["Fouls per 90"].fillna(0)*1.0 + df["Yellow cards per 90"].fillna(0)*2.0 + df["Red cards per 90"].fillna(0)*4.0)
        df["Discipline"] = -_zscore(penalty)

    # Aux: Non-penalty goals per 90
    if all(c in df.columns for c in ["Goals","Penalties taken","Minutes played"]):
        with np.errstate(divide="ignore", invalid="ignore"):
            df["Non-penalty goals per 90"] = (df["Goals"].fillna(0) - df["Penalties taken"].fillna(0)) / (df["Minutes played"].replace(0, np.nan)/90)

    # Poaching
    poach_w = {"npxG per Shot": .35, "Goal conversion, %": .30, "Touches in box per 90": -.25,
               "Received passes per 90": -.20, "Non-penalty goals per 90": .40}
    parts = []
    for c, w in poach_w.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if "%" in c and s.max() > 1: s = s/100.0
            parts.append(_zscore(s) * w)
    if parts: df["Poaching"] = sum(parts)

    # Finishing
    fin_w = {"Goal conversion, %": .35, "Non-penalty goals per 90": .30, "Shots on target, %": .25, "G-xG": .15, "npxG per Shot": .10}
    parts = []
    for c, w in fin_w.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if "%" in c and s.max() > 1: s = s/100.0
            parts.append(_zscore(s) * w)
    if parts: df["Finishing"] = sum(parts)

    # Aerial Threat
    aer_w = {"Head goals per 90": .35, "Aerial duels per 90": .20, "Aerial duels won, %": .20}
    parts = []
    for c, w in aer_w.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if "%" in c and s.max() > 1: s = s/100.0
            parts.append(_zscore(s) * w)
    if parts: df["Aerial Threat"] = sum(parts)

    # Passing Quality
    pq_pairs = {
        "Passes to final third per 90": ("Accurate passes to final third, %", 0.35),
        "Forward passes per 90": ("Accurate forward passes, %", 0.30),
        "Long passes per 90": ("Accurate long passes, %", 0.15),
        "Lateral passes per 90": ("Accurate lateral passes, %", 0.10),
        "Back passes per 90": ("Accurate back passes, %", 0.05),
        "Passes per 90": ("Accurate passes, %", 0.05),
    }
    parts = []; total_w = 0.0
    for vol_col, (acc_col, w) in pq_pairs.items():
        if vol_col in df.columns and acc_col in df.columns:
            vol_z = _zscore(df[vol_col].fillna(0))
            acc = df[acc_col].fillna(0)
            if acc.max() > 1: acc = acc/100.0
            acc_z = _zscore(acc)
            parts.append((0.3*vol_z + 0.7*acc_z) * w); total_w += w
    if parts and total_w > 0:
        df["Passing Quality"] = sum(parts) / total_w

    # Aerial Defence
    ad_w = {"Aerial duels per 90": .35, "Aerial duels won, %": .40, "PAdj Interceptions": .10, "Shots blocked per 90": .10}
    parts = []
    for c, w in ad_w.items():
        if c in df.columns:
            s = df[c].fillna(0)
            if "%" in c and s.max() > 1: s = s/100.0
            parts.append(_zscore(s) * w)
    if parts: df["Aerial Defence"] = sum(parts)

    # Normalize to [0,100]
    main_derived = [
        "npxG","npxG per 90","npxG per Shot","Box Threat","xG Buildup","Creativity","Progression",
        "Defence","Involvement","Discipline","G-xG","Poaching","Finishing","Aerial Threat","Passing Quality","Aerial Defence"
    ]
    for m in main_derived:
        if m in df.columns:
            s = pd.to_numeric(df[m], errors="coerce")
            if s.notna().sum() > 1:
                lo, hi = np.nanmin(s), np.nanmax(s)
                if np.isfinite(lo) and np.isfinite(hi) and hi != lo:
                    df[m] = (s - lo) / (hi - lo) * 100.0
                else:
                    df[m] = 50.0
            else:
                df[m] = 50.0

    # Invert negatives if present
    for m in ["Conceded goals per 90","xG against per 90","Fouls per 90","Yellow cards per 90","Red cards per 90"]:
        if m in df.columns:
            s = pd.to_numeric(df[m], errors="coerce")
            if s.notna().sum() > 1:
                hi, lo = np.nanmax(s), np.nanmin(s)
                if hi != lo: df[m] = (hi - s) / (hi - lo) * 100.0
                else: df[m] = 50.0
    return df

# ===================== OUR OFF/DEF COMPOSITES =====================
def compute_offensive_production(df: pd.DataFrame, weights: dict[str, float]) -> pd.Series:
    df = _ensure_cols(df, OFFENSIVE_COMPONENTS)
    s = pd.Series(0.0, index=df.index)
    for col, w in weights.items():
        if col in df.columns:
            s = s + _safe_s(df, col) * float(w)
    return s

def compute_defensive_production(df: pd.DataFrame) -> pd.Series:
    df = _ensure_cols(df, DEF_COLS)
    def_eff = _safe_s(df, "Defensive duels per 90") * (_safe_s(df, "Defensive duels won, %")/100.0)
    aer_eff = _safe_s(df, "Aerial duels per 90") * (_safe_s(df, "Aerial duels won, %")/100.0)
    parts = [
        _safe_s(df, "Successful defensive actions per 90"),
        def_eff, aer_eff,
        _safe_s(df, "PAdj Sliding tackles"),
        _safe_s(df, "PAdj Interceptions"),
        _safe_s(df, "Shots blocked per 90"),
    ]
    total = parts[0]
    for p in parts[1:]:
        total = total + p
    return total

def compute_composite_metrics(df: pd.DataFrame, off_weights: dict[str, float]) -> pd.DataFrame:
    df = df.copy()
    df = compute_derived_metrics(df)

    df = _ensure_cols(df, list(PHYS_COLS.values()))
    df["Prod_Ofensiva_p90"] = compute_offensive_production(df, off_weights)
    df["Prod_Defensiva_p90"] = compute_defensive_production(df)

    def pick(*cands):
        for c in cands:
            if c in df.columns: return c
        return None
    dist_col = pick("Distance P90","distance_total_m","Total distance per 90","Total Distance per 90 (m)")
    run_col  = pick("Running Distance P90","running_m")
    hi_col   = pick("HI Distance P90","hi_m")
    hsr_cnt  = pick("HSR Count P90","hsr_cnt")
    spr_cnt  = pick("Sprint Count P90","sprint_cnt")
    ex_hsr   = pick("Explosive Acceleration to HSR Count P90","expl_hsr_cnt")
    ex_spr   = pick("Explosive Acceleration to Sprint Count P90","expl_sprint_cnt")

    dist_km = None
    if dist_col:
        series = pd.to_numeric(df[dist_col], errors="coerce").fillna(0)
        dist_km = np.where(series.abs().max() > 2000, series/1000.0, series)

    hi_run_km = None
    if hi_col or run_col:
        s_hi = pd.to_numeric(df[hi_col], errors="coerce").fillna(0) if hi_col else 0.0
        s_run = pd.to_numeric(df[run_col], errors="coerce").fillna(0) if run_col else 0.0
        s_both = s_hi + s_run
        hi_run_km = np.where(s_both.abs().max() > 2000, s_both/1000.0, s_both)

    exp_events = None
    cnts = []
    for c in [hsr_cnt, spr_cnt, ex_hsr, ex_spr]:
        if c: cnts.append(pd.to_numeric(df[c], errors="coerce").fillna(0))
    if cnts: exp_events = sum(cnts)

    def safe_ratio(num, den):
        if num is None or den is None: return np.nan
        den = np.asarray(den); den = np.where(den==0, np.nan, den)
        return num / den

    if dist_km is not None:
        df["Work Rate Offensive"] = safe_ratio(df["Prod_Ofensiva_p90"], dist_km)
        df["Work Rate Defensive"] = safe_ratio(df["Prod_Defensiva_p90"], dist_km)
    if hi_run_km is not None:
        df["Offensive Intensity"] = safe_ratio(df["Prod_Ofensiva_p90"], hi_run_km)
        df["Defensive Intensity"] = safe_ratio(df["Prod_Defensiva_p90"], hi_run_km)
    if exp_events is not None:
        df["Offensive Explosion"] = safe_ratio(df["Prod_Ofensiva_p90"], exp_events)
        df["Defensive Explosion"] = safe_ratio(df["Prod_Defensiva_p90"], exp_events)

    if "Minutes played" not in df.columns:
        raise ValueError("A coluna 'Minutes played' não foi encontrada no arquivo enviado. Certifique-se de manter exatamente esse nome.")
    return df

# ===================== RADAR / UI HELPERS =====================
def _bounds_from_df(df: pd.DataFrame, metrics: list[str]):
    lowers, uppers = [], []
    for m in metrics:
        s = pd.to_numeric(df[m], errors="coerce")
        if m in NEGATE_METRICS: s = -s
        lo = np.nanpercentile(s, 5); hi = np.nanpercentile(s, 95)
        if not np.isfinite(lo) or not np.isfinite(hi) or lo == hi:
            lo = np.nanmin(s); hi = np.nanmax(s)
        if lo == hi: hi = lo + 1e-6
        lowers.append(float(lo)); uppers.append(float(hi))
    return lowers, uppers

def _values_for_player(row: pd.Series, metrics: list[str]):
    vals = []
    for m in metrics:
        v = pd.to_numeric(row.get(m, np.nan), errors="coerce")
        if m in NEGATE_METRICS and pd.notna(v): v = -v
        vals.append(float(v) if pd.notna(v) else np.nan)
    return vals

def _player_label(row: pd.Series) -> str:
    name = str(row.get("Player", ""))
    team = str(row.get("Team", "")) if "Team" in row.index and pd.notna(row.get("Team")) else ""
    pos  = str(row.get("Position", "")) if "Position" in row.index and pd.notna(row.get("Position")) else ""
    minutes = None
    for mcol in ["Minutes played","Minutes","minutes","Time played","Minuti","Minutos","Min"]:
        if mcol in row.index and pd.notna(row.get(mcol)):
            try:
                minutes = int(float(row[mcol]))
            except Exception:
                minutes = row[mcol]
            break
    tail = f" — {team}" if team else ""
    if pos: tail += f" | {pos}"
    if minutes is not None: tail += f" | {minutes} min"
    return name + tail

def plot_radar(df: pd.DataFrame, player_a: str, player_b: str | None, metrics: list[str],
               color_a: str, color_b: str = "#E76F51"):
    if not metrics:
        st.warning("Select at least 3 metrics for the radar.")
        return
    metrics = metrics[:16]
    row_a = df[df["Player"] == player_a].iloc[0]
    row_b = df[df["Player"] == player_b].iloc[0] if player_b else None
    lowers, uppers = _bounds_from_df(df, metrics)
    radar = Radar(metrics, lowers, uppers, num_rings=4)

    v_a = _values_for_player(row_a, metrics)
    v_b = _values_for_player(row_b, metrics) if row_b is not None else None

    fig, ax = plt.subplots(figsize=(8, 8))
    radar.setup_axis(ax=ax)
    radar.draw_circles(ax=ax, facecolor="#f3f3f3", edgecolor="#c9c9c9", alpha=0.18)
    try:
        radar.spoke(ax=ax, color="#c9c9c9", linestyle="--", alpha=0.18)
    except Exception:
        pass
    radar.draw_radar(v_a, ax=ax, kwargs_radar={"facecolor": color_a+"33", "edgecolor": color_a, "linewidth": 2})
    if v_b is not None:
        radar.draw_radar(v_b, ax=ax, kwargs_radar={"facecolor": color_b+"33", "edgecolor": color_b, "linewidth": 2})
    radar.draw_range_labels(ax=ax, fontsize=9)
    radar.draw_param_labels(ax=ax, fontsize=10)

    title_a = _player_label(row_a)
    title = title_a if row_b is None else f"{title_a} vs {_player_label(row_b)}"
    ax.set_title(title, fontsize=14, pad=20)

    st.pyplot(fig, use_container_width=True)

# ===================== Ranking bars helpers =====================
def _metric_rank_info(dfin: pd.DataFrame, metric: str, player_name: str):
    # Barra baseada em ranking: 1 = melhor => barra cheia (norm = 1.0)
    s = pd.to_numeric(dfin[metric], errors="coerce")
    mask = s.notna()
    s = s[mask]
    d = dfin.loc[mask]
    total = int(s.shape[0]) if s.shape[0] > 0 else 0
    if total == 0:
        return {"rank": None, "total": 0, "value": np.nan, "norm": 0.0, "ascending": False}
    ascending = metric in NEGATE_METRICS  # se True, menor é melhor

    # Rank ordinal (1 = melhor considerando ascending)
    r = s.rank(ascending=ascending, method="min")
    try:
        player_idx = d.index[d["Player"] == player_name]
        if player_idx.empty:
            return {"rank": None, "total": total, "value": np.nan, "norm": 0.0, "ascending": ascending}
        val = float(s.loc[player_idx[0]]) if player_idx[0] in s.index else np.nan
        rk = int(r.loc[player_idx[0]]) if player_idx[0] in r.index and not np.isnan(r.loc[player_idx[0]]) else None
    except Exception:
        val, rk = np.nan, None

    # Converte rank -> [0,1]: 1.0 para rank 1; 0.0 para rank = total
    if rk is None:
        norm = 0.0
    elif total <= 1:
        norm = 1.0
    else:
        norm = 1.0 - (float(rk - 1) / float(total - 1))

    norm = float(np.clip(norm, 0.0, 1.0))
    return {"rank": rk, "total": total, "value": val, "norm": norm, "ascending": ascending}


def _metric_percentile_info(dfin: pd.DataFrame, metric: str, player_name: str):
    if metric not in dfin.columns:
        return {"percentile": np.nan, "total": 0, "value": np.nan}

    s = pd.to_numeric(dfin[metric], errors="coerce")
    mask = s.notna()
    total = int(mask.sum())
    if total == 0:
        return {"percentile": np.nan, "total": 0, "value": np.nan}

    df_valid = dfin.loc[mask]
    if "Player" not in df_valid.columns:
        return {"percentile": np.nan, "total": total, "value": np.nan}

    player_idx = df_valid.index[df_valid["Player"] == player_name]
    if player_idx.empty:
        return {"percentile": np.nan, "total": total, "value": np.nan}

    idx = player_idx[0]
    val = float(s.loc[idx]) if idx in s.index and pd.notna(s.loc[idx]) else np.nan

    ascending = metric in NEGATE_METRICS
    if total <= 1:
        pct = 100.0
    else:
        ranks = s[mask].rank(ascending=ascending, method="average")
        if idx in ranks.index and not np.isnan(ranks.loc[idx]):
            pct = 100.0 * (total - ranks.loc[idx]) / (total - 1)
        else:
            pct = np.nan

    if pd.isna(pct):
        pct_val = np.nan
    else:
        pct_val = float(np.clip(pct, 0.0, 100.0))

    return {"percentile": pct_val, "total": total, "value": val}


def _format_metric_value(metric: str, value: float | None) -> str:
    if value is None or not np.isfinite(value):
        return "—"
    if "%" in metric:
        val = float(value)
        if abs(val) <= 1.2:
            val *= 100.0
        return f"{val:.1f}%"
    abs_val = abs(value)
    if abs_val >= 100:
        return f"{value:.0f}"
    if abs_val >= 10:
        return f"{value:.1f}"
    if abs_val >= 1:
        return f"{value:.2f}"
    return f"{value:.3f}"


def _prepare_metric_bar_context(dfin: pd.DataFrame, metric: str, player_names: list[str]):
    transform = -1.0 if metric in NEGATE_METRICS else 1.0
    series = pd.to_numeric(dfin.get(metric, pd.Series(dtype=float)), errors="coerce")
    series = series.replace([np.inf, -np.inf], np.nan).dropna()
    plot_series = series * transform

    if plot_series.empty:
        base_lo, base_hi = -0.5, 0.5
        q1 = median = q3 = np.nan
    else:
        try:
            base_lo = float(np.nanpercentile(plot_series, 5))
            base_hi = float(np.nanpercentile(plot_series, 95))
        except Exception:
            base_lo, base_hi = float(np.nanmin(plot_series)), float(np.nanmax(plot_series))
        if not np.isfinite(base_lo) or not np.isfinite(base_hi):
            base_lo, base_hi = -0.5, 0.5
        if base_lo == base_hi:
            base_hi = base_lo + 1.0 if base_lo == 0 else base_lo + abs(base_lo) * 0.1
        try:
            q1 = float(np.nanpercentile(plot_series, 25))
            median = float(np.nanpercentile(plot_series, 50))
            q3 = float(np.nanpercentile(plot_series, 75))
        except Exception:
            q1 = median = q3 = np.nan

    axis_lo, axis_hi = base_lo, base_hi
    player_infos = []
    for idx, name in enumerate(player_names):
        info = _metric_rank_info(dfin, metric, name)
        info["player"] = name
        value = info.get("value")
        if value is None or not np.isfinite(value):
            info["plot_value"] = np.nan
        else:
            plot_val = float(value) * transform
            info["plot_value"] = plot_val
            if axis_lo is None or not np.isfinite(axis_lo):
                axis_lo = plot_val
            else:
                axis_lo = min(axis_lo, plot_val)
            if axis_hi is None or not np.isfinite(axis_hi):
                axis_hi = plot_val
            else:
                axis_hi = max(axis_hi, plot_val)
        player_infos.append(info)

    if axis_lo is None or axis_hi is None or not np.isfinite(axis_lo) or not np.isfinite(axis_hi):
        axis_lo, axis_hi = base_lo, base_hi
    if axis_lo == axis_hi:
        axis_hi = axis_lo + 1.0 if axis_lo == 0 else axis_lo + abs(axis_lo) * 0.1

    stats = {
        "transform": transform,
        "range_lo": base_lo,
        "range_hi": base_hi,
        "q1": q1,
        "median": median,
        "q3": q3,
        "axis_lo": axis_lo,
        "axis_hi": axis_hi,
    }
    return stats, player_infos


def _draw_metric_bar_axis(ax, metric: str, stats: dict, player_infos: list[dict], colors: list[str]):
    transform = stats.get("transform", 1.0)
    axis_lo = stats.get("axis_lo", -0.5)
    axis_hi = stats.get("axis_hi", 0.5)
    if not np.isfinite(axis_lo):
        axis_lo = -0.5
    if not np.isfinite(axis_hi):
        axis_hi = 0.5
    if axis_lo == axis_hi:
        axis_hi = axis_lo + 1.0 if axis_lo == 0 else axis_lo + abs(axis_lo) * 0.1
    span = axis_hi - axis_lo
    margin = span * 0.06 if span != 0 else 1.0
    ax.set_xlim(axis_lo - margin, axis_hi + margin)
    ax.set_ylim(0, 1)
    ax.set_yticks([])
    ax.set_axisbelow(True)

    range_lo = stats.get("range_lo")
    range_hi = stats.get("range_hi")
    if np.isfinite(range_lo) and np.isfinite(range_hi):
        lo_bg, hi_bg = sorted([range_lo, range_hi])
        ax.axvspan(lo_bg, hi_bg, color="#e2e8f0", alpha=0.6, zorder=0)

    for val, style in [(stats.get("q1"), ":"), (stats.get("median"), "--"), (stats.get("q3"), ":")]:
        if np.isfinite(val):
            ax.axvline(val, color="#94a3b8", linestyle=style, linewidth=0.7, zorder=1)

    ax.grid(axis="x", linestyle=":", linewidth=0.4, alpha=0.5, color="#cbd5e1")

    valid_infos = [(idx, info) for idx, info in enumerate(player_infos) if np.isfinite(info.get("plot_value", np.nan))]
    missing_infos = [info for info in player_infos if not np.isfinite(info.get("plot_value", np.nan))]

    if len(valid_infos) == 1:
        positions = [0.55]
    elif len(valid_infos) > 1:
        positions = np.linspace(0.7, 0.3, len(valid_infos))
    else:
        positions = []

    baseline = range_lo if np.isfinite(range_lo) else axis_lo
    if not np.isfinite(baseline):
        baseline = axis_lo

    x_min, x_max = ax.get_xlim()
    text_offset = span * 0.02 if span != 0 else 0.05
    for pos, (idx, info) in zip(positions, valid_infos):
        color = colors[idx] if idx < len(colors) else "#1d4ed8"
        plot_val = info["plot_value"]
        width = plot_val - baseline
        ax.barh(pos, width, left=baseline, height=0.18, color=color, alpha=0.85, zorder=3)

        actual_value = info.get("value")
        value_txt = _format_metric_value(metric, actual_value)
        rank = info.get("rank")
        total = info.get("total")
        rank_txt = f" (#{rank}/{total})" if rank is not None and total else ""
        label = f"{info['player']}: {value_txt}{rank_txt}"

        if width >= 0:
            text_x = plot_val + text_offset
            ha = "left"
        else:
            text_x = plot_val - text_offset
            ha = "right"
        text_x = float(np.clip(text_x, axis_lo - margin, axis_hi + margin))
        ax.text(text_x, pos, label, va="center", ha=ha, fontsize=8, color="#0f172a", zorder=4)

    for miss_idx, info in enumerate(missing_infos):
        ax.text(
            0.02,
            0.08 + miss_idx * 0.12,
            f"{info['player']}: sem dados",
            transform=ax.transAxes,
            ha="left",
            va="bottom",
            fontsize=7,
            color="#94a3b8",
        )

    arrow_note = " (menor é melhor)" if transform == -1 else ""
    ax.set_title(f"{metric}{arrow_note}", fontsize=9, pad=4, loc="left", color="#0f172a")

    ax.xaxis.set_major_locator(MaxNLocator(nbins=4, prune="both"))
    ax.xaxis.set_major_formatter(FuncFormatter(lambda x, _: _format_metric_value(metric, x * transform)))
    ax.tick_params(axis="x", labelsize=7, colors="#475569", pad=1)

    for spine in ax.spines.values():
        spine.set_visible(False)

def render_metric_rank_bars(dfin: pd.DataFrame, player_a: str, metrics: list[str], player_b: str | None = None):
    if not metrics:
        return
    st.markdown("### 📊 Ranking por métrica")
    st.caption("Barra indica desempenho relativo; rótulo mostra a posição no ranking (1 = melhor).")

    def _render_for(player_name: str, header: str):
        st.markdown(f"**{header}:** {player_name}")
        cols_per_row = 3
        for i, m in enumerate(metrics):
            if i % cols_per_row == 0:
                cols = st.columns(cols_per_row)
            c = cols[i % cols_per_row]
            with c:
                info = _metric_rank_info(dfin, m, player_name)
                rk, tot, val, norm = info["rank"], info["total"], info["value"], info["norm"]
                label = f"{m} — {rk}/{tot}" if rk is not None else f"{m} — n/a"
                fig, ax = plt.subplots(figsize=(4, 0.6))
                ax.barh([0], [norm])
                ax.set_xlim(0, 1)
                ax.set_yticks([])
                ax.set_xticks([0, 0.5, 1])
                ax.set_xticklabels(["0%","50%","100%"], fontsize=7)
                ax.set_title(label, fontsize=9, pad=2)
                for spine in ["top","right","left"]:
                    ax.spines[spine].set_visible(False)
                st.pyplot(fig, use_container_width=True)

    _render_for(player_a, "Jogador A")
    if player_b:
        _render_for(player_b, "Jogador B")

# ======= Build a single PNG that includes Radar + Ranking/Percentile Bars =======
def make_radar_bars_png(
    df: pd.DataFrame,
    player_a: str,
    player_b: str | None,
    metrics: list[str],
    color_a: str,
    color_b: str = "#E76F51",
    bar_mode: str = "rank",
) -> io.BytesIO:
    metrics = [m for m in (metrics or []) if m in df.columns][:16]
    if not metrics:
        raise ValueError("É necessário informar métricas válidas para gerar o radar.")

    bar_mode = (bar_mode or "rank").lower()
    if bar_mode not in {"rank", "percentile"}:
        raise ValueError("bar_mode deve ser 'rank' ou 'percentile'.")

    lowers, uppers = _bounds_from_df(df, metrics)
    radar = Radar(metrics, lowers, uppers, num_rings=4)

    row_a = df[df["Player"] == player_a].iloc[0]
    v_a = _values_for_player(row_a, metrics)

    row_b = None
    v_b = None
    if player_b:
        row_b = df[df["Player"] == player_b].iloc[0]
        v_b = _values_for_player(row_b, metrics)

    cols_per_row = 3
    rows_per_player = max(1, math.ceil(len(metrics) / cols_per_row))
    bar_blocks = 1 + (1 if player_b else 0)
    total_bar_rows = rows_per_player * bar_blocks

    base_height = 9.4
    fig = plt.figure(figsize=(8.3, base_height + total_bar_rows * 0.38))
    height_ratios = [4.4, 3.4] + [0.4] * total_bar_rows if total_bar_rows else [4.4, 3.4]
    gs = GridSpec(
        nrows=2 + total_bar_rows,
        ncols=3,
        figure=fig,
        height_ratios=height_ratios,
    )

    # Radar spans first 2 rows
    ax_radar = fig.add_subplot(gs[0:2, :])
    radar.setup_axis(ax=ax_radar)
    radar.draw_circles(
        ax=ax_radar,
        facecolor="#f8fafc",
        edgecolor="#cbd5f5",
        alpha=0.22,
    )
    try:
        radar.spoke(ax=ax_radar, color="#c9c9c9", linestyle="--", alpha=0.18)
    except Exception:
        pass
    radar.draw_radar(v_a, ax=ax_radar, kwargs_radar={"facecolor": color_a+"33", "edgecolor": color_a, "linewidth": 2})
    if v_b is not None:
        radar.draw_radar(v_b, ax=ax_radar, kwargs_radar={"facecolor": color_b+"33", "edgecolor": color_b, "linewidth": 2})
    radar.draw_range_labels(ax=ax_radar, fontsize=9)
    radar.draw_param_labels(ax=ax_radar, fontsize=11)

    title_a = _player_label(row_a)
    title = title_a if row_b is None else f"{title_a} vs {_player_label(row_b)}"
    ax_radar.set_title(title, fontsize=20, weight="bold", pad=18)

    scale_max = 100.0 if bar_mode == "percentile" else 1.0
    if bar_mode == "percentile":
        xticks = [0, 10, 50, 90, 100]
        xticklabels = [f"{tick:.0f}%" for tick in xticks]
    else:
        xticks = [0, 0.5, 1]
        xticklabels = ["0%", "50%", "100%"]

    if bar_mode == "percentile":
        max_radius = radar.center_circle_radius + radar.ring_width * radar.num_rings
        span = max_radius - radar.center_circle_radius
        for anchor in (10, 50, 90):
            radius = radar.center_circle_radius + (anchor / 100.0) * span
            anchor_circle = plt.Circle(
                (0, 0),
                radius,
                transform=ax_radar.transData,
                fill=False,
                linestyle="--",
                linewidth=0.6,
                color="#94a3b8",
                alpha=0.45,
            )
            ax_radar.add_patch(anchor_circle)
            ax_radar.text(
                radius + 0.1,
                0,
                f"{anchor}",
                fontsize=8,
                color="#475569",
                va="center",
                ha="left",
            )

    # Bar blocks (player A then optional player B)
    def _draw_bar_block(start_row: int, player_name: str):
        for i, m in enumerate(metrics):
            r = start_row + (i // cols_per_row)
            c = i % cols_per_row
            ax = fig.add_subplot(gs[2 + r, c])

            if bar_mode == "percentile":
                info = _metric_percentile_info(df, m, player_name)
                pct = info.get("percentile")
                bar_value = float(np.clip(pct, 0.0, 100.0)) if pd.notna(pct) else 0.0
                value_note = _format_metric_value(m, info.get("value"))
                value_suffix = f" ({value_note})" if value_note != "—" else ""
                label = f"{m} — {bar_value:.0f} pct{value_suffix}" if pd.notna(pct) else f"{m} — n/a{value_suffix}"
            else:
                info = _metric_rank_info(df, m, player_name)
                norm = info.get("norm")
                bar_value = float(np.clip(norm, 0.0, 1.0)) if norm is not None else 0.0
                value_note = _format_metric_value(m, info.get("value"))
                value_suffix = f" ({value_note})" if value_note != "—" else ""
                rk, tot = info.get("rank"), info.get("total")
                label = f"{m} — {rk}/{tot}{value_suffix}" if rk is not None else f"{m} — n/a{value_suffix}"

                # Converte para escala de exibição
                bar_value *= scale_max

            bar_color = color_a if player_name == player_a else color_b
            ax.barh([0], [bar_value], color=bar_color)
            ax.set_xlim(0, scale_max)
            ax.set_yticks([])
            ax.set_xticks(xticks)
            ax.set_xticklabels(xticklabels, fontsize=7)
            if bar_mode == "percentile":
                for anchor in (10, 50, 90):
                    ax.axvline(anchor, color="#cbd5f5", linestyle="--", linewidth=0.8, alpha=0.6, zorder=0)
                ax.scatter(
                    [10, 50, 90],
                    [0, 0, 0],
                    color="#2563EB",
                    s=10,
                    marker="|",
                    linewidths=0.8,
                    zorder=3,
                )
            ax.set_title(label, fontsize=9, pad=2, loc="left")
            for spine in ["top", "right", "left"]:
                ax.spines[spine].set_visible(False)

    _draw_bar_block(start_row=0, player_name=player_a)
    if player_b:
        _draw_bar_block(start_row=rows_per_player, player_name=player_b)

    buf = io.BytesIO()
    fig.tight_layout(pad=1.1)
    fig.savefig(buf, format="png", dpi=220, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf



def build_player_report_docx(
    df: pd.DataFrame,
    player_name: str,
    metrics: list[str],
    color_a: str,
    color_b: str = "#E76F51",
    player_photo: bytes | io.BytesIO | None = None,
    team_logo: bytes | io.BytesIO | None = None,
    report_title: str | None = None,
    report_subtitle: str | None = None,
    competition_level_override: str | None = None,
    prepared_by: str | None = None,
    confidentiality_note: str | None = None,
    generated_on: str | None = None,
    report_date: str | None = None,
) -> io.BytesIO:
    if "Player" not in df.columns:
        raise ValueError("DataFrame must contain the 'Player' column.")

    metrics = [m for m in (metrics or []) if m in df.columns][:16]
    if len(metrics) < 3:
        raise ValueError("Please select at least three valid metrics for the DOCX report.")

    player_rows = df[df["Player"] == player_name]
    if player_rows.empty:
        raise ValueError(f"Player '{player_name}' was not found in the dataset.")

    row = player_rows.iloc[0]

    def _to_stream(data: bytes | io.BytesIO | None) -> io.BytesIO | None:
        if data is None:
            return None
        if isinstance(data, io.BytesIO):
            data.seek(0)
            return data
        if isinstance(data, (bytes, bytearray)):
            return io.BytesIO(data)
        if hasattr(data, "read"):
            current_pos = data.tell() if hasattr(data, "tell") else None
            content = data.read()
            if current_pos is not None:
                data.seek(current_pos)
            return io.BytesIO(content)
        raise TypeError("Images must be provided as bytes or in-memory files (BytesIO).")

    photo_stream = _to_stream(player_photo)
    logo_stream = _to_stream(team_logo)

    def _clone_stream(stream: io.BytesIO | None) -> io.BytesIO | None:
        if stream is None:
            return None
        try:
            stream.seek(0)
            data = stream.read()
            stream.seek(0)
        except Exception:
            return None
        clone = io.BytesIO(data)
        clone.seek(0)
        return clone

    accent_hex = "2563EB"
    accent_color = f"#{accent_hex}"
    neutral_border_hex = "D1D5DB"
    neutral_fill = "F3F4F6"
    zebra_fill = "FAFAFA"
    text_primary = RGBColor(31, 41, 55)
    muted_text = RGBColor(107, 114, 128)

    primary_title = report_title or REPORT_TITLE
    subtitle = report_subtitle or REPORT_SUBTITLE
    athlete_name = player_name or PLAYER_NAME_DEFAULT
    default_confidentiality_note = (
        confidentiality_note or REPORT_CONFIDENTIALITY_NOTE
    )
    if generated_on is None:
        generated_on = datetime.now().strftime("%B %d, %Y")

    def _clean(value):
        if value is None:
            return None
        if isinstance(value, str):
            txt = value.strip()
            return txt or None
        if pd.isna(value):
            return None
        return value

    confidentiality_note = _clean(default_confidentiality_note) or REPORT_CONFIDENTIALITY_NOTE
    prepared_by = _clean(prepared_by) or REPORT_AUTHOR
    generated_on = _clean(generated_on) or datetime.now().strftime("%B %d, %Y")
    report_date = _clean(report_date)

    player_club = _clean(row.get("Team")) or PLAYER_CLUB_DEFAULT
    player_position = _clean(row.get("Position")) or PLAYER_POSITION_DEFAULT
    player_age = _player_age(row) or PLAYER_AGE_DEFAULT
    competition_level = (
        _clean(competition_level_override)
        or _clean(row.get("League"))
        or COMPETITION_LEVEL_DEFAULT
    )
    if "Player" in df.columns:
        sample_size = int(pd.Series(df["Player"]).dropna().nunique())
    else:
        sample_size = len(df)
    if not sample_size:
        sample_size = SAMPLE_SIZE_DEFAULT

    def _descriptor_for_metric(metric_name: str, pct_value: float, raw_value):
        descriptor = f"{pct_value:.0f}th pct"
        formatted = _format_metric_value(metric_name, raw_value)
        if formatted and formatted != "—":
            descriptor = f"{descriptor} ({formatted})"
        return descriptor

    insight_pool: list[dict[str, object]] = []

    def _metric_median(metric_name: str) -> float | None:
        if metric_name not in df.columns:
            return None
        series = pd.to_numeric(df[metric_name], errors="coerce")
        series = series.dropna()
        if series.empty:
            return None
        return float(series.median())

    for metric in metrics:
        info = _metric_percentile_info(df, metric, player_name)
        pct_value = info.get("percentile")
        if pd.notna(pct_value):
            pct_float = float(pct_value)
            descriptor = _descriptor_for_metric(metric, pct_float, info.get("value"))
            insight_pool.append(
                {
                    "metric": metric,
                    "pct": pct_float,
                    "descriptor": descriptor,
                    "source": "radar",
                    "value": info.get("value"),
                    "median": _metric_median(metric),
                }
            )

    candidate_additional_metrics = [
        col
        for col in df.columns
        if col not in metrics
        and col not in ID_COLS_CANDIDATES
        and pd.api.types.is_numeric_dtype(df[col])
    ]

    for metric in candidate_additional_metrics:
        info = _metric_percentile_info(df, metric, player_name)
        pct_value = info.get("percentile")
        if not pd.notna(pct_value):
            continue
        pct_float = float(pct_value)
        if pct_float >= 70.0 or pct_float <= 40.0:
            descriptor = _descriptor_for_metric(metric, pct_float, info.get("value"))
            insight_pool.append(
                {
                    "metric": metric,
                    "pct": pct_float,
                    "descriptor": descriptor,
                    "source": "extra",
                    "value": info.get("value"),
                    "median": _metric_median(metric),
                }
            )

    if not insight_pool:
        for item in STANDOUTS_DEFAULT + DEVELOPMENT_AREAS_DEFAULT:
            insight_pool.append(
                {
                    "metric": item.get("metric"),
                    "pct": item.get("pct"),
                    "descriptor": item.get("descriptor"),
                    "source": "default",
                    "value": item.get("value"),
                    "median": item.get("median"),
                }
            )

    def _select_insights(
        pool: list[dict[str, object]],
        predicate,
        *,
        limit: int,
        reverse: bool,
    ) -> list[dict[str, object]]:
        filtered: list[dict[str, object]] = [
            item
            for item in pool
            if predicate(item) and item.get("descriptor")
        ]
        filtered.sort(key=lambda item: item.get("pct", 0.0) or 0.0, reverse=reverse)

        unique: list[dict[str, object]] = []
        seen: set[str] = set()
        for item in filtered:
            metric_name = str(item.get("metric"))
            if metric_name in seen:
                continue
            unique.append(item)
            seen.add(metric_name)
            if len(unique) == limit:
                break

        extra_candidates = [
            item
            for item in filtered
            if item.get("source") == "extra"
            and str(item.get("metric")) not in seen
        ]
        if extra_candidates and not any(it.get("source") == "extra" for it in unique):
            replacement = extra_candidates[0]
            if unique and len(unique) >= limit:
                unique[-1] = replacement
            else:
                unique.append(replacement)

        return unique[:limit]

    standouts = _select_insights(
        insight_pool,
        lambda item: item.get("pct") is not None and item.get("pct") >= 70.0,
        limit=5,
        reverse=True,
    )
    if not standouts:
        standouts = list(STANDOUTS_DEFAULT)

    development = _select_insights(
        insight_pool,
        lambda item: item.get("pct") is not None and item.get("pct") <= 40.0,
        limit=5,
        reverse=False,
    )
    if not development:
        development = list(DEVELOPMENT_AREAS_DEFAULT)

    doc = Document()

    def _configure_section(section):
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Mm(A4_WIDTH_MM)
        section.page_height = Mm(A4_HEIGHT_MM)
        margin = Mm(PAGE_MARGIN_MM)
        section.left_margin = section.right_margin = margin
        section.top_margin = section.bottom_margin = margin
        section.header_distance = Mm(HEADER_FOOTER_DISTANCE_MM)
        section.footer_distance = Mm(HEADER_FOOTER_DISTANCE_MM)

    def _ensure_style(name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
        try:
            return doc.styles[name]
        except KeyError:
            return doc.styles.add_style(name, style_type)

    normal = _ensure_style("Normal")
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = text_primary
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title_style = _ensure_style("TitleHero")
    title_style.font.name = "Calibri"
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = text_primary
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(10)
    title_style.paragraph_format.line_spacing = 1.1

    h1_style = _ensure_style("H1")
    h1_style.font.name = "Calibri"
    h1_style.font.size = Pt(20)
    h1_style.font.bold = True
    h1_style.font.color.rgb = text_primary
    h1_style.paragraph_format.space_before = Pt(4)
    h1_style.paragraph_format.space_after = Pt(6)
    h1_style.paragraph_format.line_spacing = 1.1
    h1_style.paragraph_format.keep_with_next = True
    h1_style.paragraph_format.page_break_before = False

    h2_style = _ensure_style("H2")
    h2_style.font.name = "Calibri"
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = text_primary
    h2_style.paragraph_format.space_before = Pt(2)
    h2_style.paragraph_format.space_after = Pt(2)
    h2_style.paragraph_format.keep_with_next = True

    body_style = _ensure_style("Body")
    body_style.font.name = "Calibri"
    body_style.font.size = Pt(11)
    body_style.font.color.rgb = text_primary

    small_caps_style = _ensure_style("SmallCaps")
    small_caps_style.font.name = "Calibri"
    small_caps_style.font.size = Pt(12)
    small_caps_style.font.small_caps = True
    small_caps_style.font.color.rgb = muted_text

    tag_style = _ensure_style("Tag")
    tag_style.font.name = "Calibri"
    tag_style.font.size = Pt(10)
    tag_style.font.color.rgb = text_primary

    kpi_style = _ensure_style("KPI")
    kpi_style.font.name = "Calibri"
    kpi_style.font.size = Pt(13)
    kpi_style.font.bold = True
    kpi_style.font.color.rgb = text_primary

    note_style = _ensure_style("Note")
    note_style.font.name = "Calibri"
    note_style.font.size = Pt(10.5)
    note_style.font.italic = True
    note_style.font.color.rgb = muted_text

    for section in doc.sections:
        _configure_section(section)

    section = doc.sections[-1]

    def _apply_cell_shading(cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    def _set_cell_border(cell, **kwargs):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn("w:tcBorders"))
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for edge in ("top", "left", "bottom", "right"):
            edge_data = kwargs.get(edge)
            if edge_data:
                element = tc_borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    tc_borders.append(element)
                element.set(qn("w:val"), edge_data.get("val", "single"))
                element.set(qn("w:sz"), str(edge_data.get("sz", 8)))
                element.set(qn("w:color"), edge_data.get("color", "000000"))

    def _set_cell_margins(cell, **kwargs):
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        for margin_name in ("top", "start", "bottom", "end"):
            margin_value = kwargs.get(margin_name)
            if margin_value is None:
                continue
            element = tc_mar.find(qn(f"w:{margin_name}"))
            if element is None:
                element = OxmlElement(f"w:{margin_name}")
                tc_mar.append(element)
            element.set(qn("w:w"), str(margin_value))
            element.set(qn("w:type"), "dxa")

    def _shade_paragraph(paragraph, fill: str):
        p = paragraph._p
        p_pr = p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)

    def _add_bottom_rule(paragraph, color: str = "E5E7EB", size: int = 6):
        p = paragraph._p
        p_pr = p.get_or_add_pPr()
        p_borders = p_pr.find(qn("w:pBdr"))
        if p_borders is None:
            p_borders = OxmlElement("w:pBdr")
            p_pr.append(p_borders)
        bottom = p_borders.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_borders.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)

    def _add_spacer(height_pt: float = 6.0) -> None:
        spacer = doc.add_paragraph()
        spacer.style = doc.styles["Body"]
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(height_pt)
        spacer.paragraph_format.line_spacing = 1
        spacer.add_run("")

    def _add_page_field(paragraph, instruction: str) -> None:
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instruction
        run._r.append(instr_text)

        fld_separate = OxmlElement("w:fldChar")
        fld_separate.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_separate)

        result_run = paragraph.add_run()
        result_text = OxmlElement("w:t")
        result_text.text = " "
        result_run._r.append(result_text)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        result_run._r.append(fld_end)

    def _force_update_fields(document: Document) -> None:
        settings_element = document.settings.element
        update_fields = settings_element.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings_element.append(update_fields)
        update_fields.set(qn("w:val"), "true")

    header = section.header
    header.is_linked_to_previous = False
    while header.paragraphs:
        p = header.paragraphs[0]._p
        p.getparent().remove(p)

    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_width_cm = (
        section.page_width.cm
        - section.left_margin.cm
        - section.right_margin.cm
    )
    usable_width_in = (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )

    header_table = header.add_table(rows=1, cols=1, width=usable_width)
    header_table.autofit = False
    header_cell = header_table.rows[0].cells[0]

    header_table.columns[0].width = Cm(max(usable_width_cm, 6.0))

    _set_cell_margins(header_cell, top=40, bottom=40, start=80, end=80)

    header_cell.text = ""
    header_title = header_cell.paragraphs[0]
    header_title.style = doc.styles["SmallCaps"]
    header_title.paragraph_format.space_after = Pt(0)
    header_title.text = primary_title
    header_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    header_subtitle = header_cell.add_paragraph(subtitle, style="Tag")
    header_subtitle.paragraph_format.space_before = Pt(0)
    header_subtitle.paragraph_format.space_after = Pt(0)
    header_subtitle.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header_meta = header_cell.add_paragraph(style="Tag")
    header_meta.paragraph_format.space_before = Pt(4)
    header_meta.paragraph_format.space_after = Pt(0)
    header_meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    header_meta_text = f"Generated on {generated_on}"
    header_meta.add_run(header_meta_text)

    footer = section.footer
    footer.is_linked_to_previous = False
    while footer.paragraphs:
        p = footer.paragraphs[0]._p
        p.getparent().remove(p)

    footer_table = footer.add_table(rows=1, cols=3, width=usable_width)
    footer_table.autofit = False
    widths = [Inches(2.8), Inches(2.8), Inches(0.9)]
    for idx, col in enumerate(footer_table.columns):
        col.width = widths[idx]

    footer_left, footer_center, footer_right = footer_table.rows[0].cells
    footer_left_paragraph = footer_left.paragraphs[0]
    footer_left_paragraph.style = doc.styles["Note"]
    footer_left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer_left_paragraph.text = ""

    footer_center_paragraph = footer_center.paragraphs[0]
    footer_center_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_center_paragraph.style = doc.styles["Body"]
    footer_center_paragraph.runs.clear()
    footer_center_paragraph.add_run("Page ")
    _add_page_field(footer_center_paragraph, "PAGE \\* Arabic")
    footer_center_paragraph.add_run(" of ")
    _add_page_field(footer_center_paragraph, "NUMPAGES \\* Arabic")


    footer_right_paragraph = footer_right.paragraphs[0]
    footer_right_paragraph.style = doc.styles["Note"]
    footer_right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_right_paragraph.text = f"Version {REPORT_VERSION}"
    footer_right_author = footer_right.add_paragraph(
        f"Prepared by {prepared_by}",
        style="Note",
    )
    footer_right_author.paragraph_format.space_before = Pt(2)
    footer_right_author.paragraph_format.space_after = Pt(0)
    footer_right_author.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    cover_table = doc.add_table(rows=1, cols=2)
    cover_table.autofit = False
    cover_table.columns[0].width = Cm(0.6)
    cover_table.columns[1].width = Cm(max(usable_width_cm - 0.6, 12.0))
    cover_row = cover_table.rows[0]
    accent_cell, hero_host_cell = cover_row.cells
    _apply_cell_shading(accent_cell, accent_hex)
    _set_cell_margins(accent_cell, top=0, bottom=0, start=0, end=0)
    accent_cell.text = ""

    hero_host_cell.text = ""
    _set_cell_margins(hero_host_cell, top=0, bottom=0, start=0, end=0)
    hero_table = hero_host_cell.add_table(rows=1, cols=3)
    hero_table.autofit = False
    hero_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hero_width_cm = max(usable_width_cm - 0.6, 12.0)
    hero_left_cm = max(3.5, hero_width_cm * 0.26)
    hero_right_cm = hero_left_cm
    hero_center_cm = max(hero_width_cm - hero_left_cm - hero_right_cm, 6.2)
    hero_table.columns[0].width = Cm(hero_left_cm)
    hero_table.columns[1].width = Cm(hero_center_cm)
    hero_table.columns[2].width = Cm(hero_right_cm)

    hero_row = hero_table.rows[0]
    for cell in hero_row.cells:
        _apply_cell_shading(cell, "F9FAFB")
        _set_cell_margins(cell, top=60, bottom=70, start=160, end=160)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    photo_cell, info_cell, crest_cell = hero_row.cells

    _set_cell_border(
        photo_cell,
        top={"sz": 12, "color": neutral_border_hex},
        bottom={"sz": 12, "color": neutral_border_hex},
        left={"sz": 12, "color": neutral_border_hex},
        right={"sz": 12, "color": neutral_border_hex},
    )
    photo_cell.text = ""
    photo_para = photo_cell.paragraphs[0]
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_para.paragraph_format.space_after = Pt(0)
    if photo_stream is not None:
        photo_stream.seek(0)
        photo_para.add_run().add_picture(photo_stream, width=Cm(hero_left_cm - 1.2))
    else:
        photo_run = photo_para.add_run("PLAYER\nPHOTO")
        photo_run.font.size = Pt(9)
        photo_run.font.bold = True
        photo_run.font.color.rgb = muted_text

    info_cell.text = ""
    title_para = info_cell.paragraphs[0]
    title_para.style = doc.styles["TitleHero"]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.text = primary_title
    title_para.paragraph_format.space_after = Pt(6)

    subtitle_para = info_cell.add_paragraph(subtitle, style="SmallCaps")
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_before = Pt(0)
    subtitle_para.paragraph_format.space_after = Pt(4)

    divider_para = info_cell.add_paragraph(" ")
    divider_para.paragraph_format.space_before = Pt(0)
    divider_para.paragraph_format.space_after = Pt(6)
    _add_bottom_rule(divider_para)

    name_para = info_cell.add_paragraph(athlete_name, style="H1")
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(6)

    strap_text = " · ".join(
        [
            value
            for value in (
                _clean(player_position) or "",
                _clean(player_club) or "",
            )
            if value
        ]
    )
    strap_line = info_cell.add_paragraph(style="Body")
    strap_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if strap_text:
        strap_line.add_run(strap_text)
    strap_line.paragraph_format.space_before = Pt(0)
    strap_line.paragraph_format.space_after = Pt(2)

    _set_cell_border(
        crest_cell,
        top={"sz": 12, "color": neutral_border_hex},
        bottom={"sz": 12, "color": neutral_border_hex},
        left={"sz": 12, "color": neutral_border_hex},
        right={"sz": 12, "color": neutral_border_hex},
    )
    crest_cell.text = ""
    crest_para = crest_cell.paragraphs[0]
    crest_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crest_para.paragraph_format.space_after = Pt(0)
    large_logo_stream = _clone_stream(logo_stream)
    if large_logo_stream is not None:
        large_logo_stream.seek(0)
        crest_para.add_run().add_picture(large_logo_stream, width=Cm(hero_right_cm - 1.2))
    else:
        crest_run = crest_para.add_run("CLUB\nCREST")
        crest_run.font.size = Pt(9)
        crest_run.font.bold = True
        crest_run.font.color.rgb = muted_text

    _add_spacer(4)

    def _add_section_heading(text: str):
        heading = doc.add_paragraph(text, style="H1")
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_bottom_rule(heading)

    _add_section_heading("Player information")

    league_parts = [
        part for part in (competition_level, report_date) if part and part != "—"
    ]
    league_display = " ".join(league_parts) if league_parts else None

    info_items = [
        ("Club", player_club),
        ("Position", player_position),
        ("Age", f"{player_age}"),
        ("League", league_display),
    ]

    info_rows = math.ceil(len(info_items) / 2)
    info_table = doc.add_table(rows=info_rows, cols=2)
    info_table.autofit = False
    column_width = Cm(max((usable_width_cm - 0.4) / 2, 6.0))
    for column in info_table.columns:
        column.width = column_width

    for index, (label, value) in enumerate(info_items):
        row_idx = index // 2
        col_idx = index % 2
        cell = info_table.rows[row_idx].cells[col_idx]
        _set_cell_border(
            cell,
            top={"sz": 14, "color": accent_hex},
            bottom={"sz": 10, "color": neutral_border_hex},
            left={"sz": 10, "color": neutral_border_hex},
            right={"sz": 10, "color": neutral_border_hex},
        )
        _set_cell_margins(cell, top=70, bottom=70, start=160, end=160)
        cell.text = ""
        band_para = cell.paragraphs[0]
        band_para.text = ""
        label_para = cell.add_paragraph(label, style="SmallCaps")
        label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_para.paragraph_format.space_before = Pt(0)
        label_para.paragraph_format.space_after = Pt(1)
        value_text = value if value not in (None, "") else "—"
        value_para = cell.add_paragraph(str(value_text), style="KPI")
        value_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        value_para.paragraph_format.space_before = Pt(0)
        value_para.paragraph_format.space_after = Pt(3)

    scope_note = doc.add_paragraph(
        "Insights reference the complete dataset, extending beyond the radar metrics for context.",
        style="Note",
    )
    scope_note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    scope_note.paragraph_format.space_before = Pt(4)
    scope_note.paragraph_format.space_after = Pt(8)

    _add_section_heading("Methodology")
    methodology_card = doc.add_table(rows=1, cols=1)
    methodology_card.autofit = False
    methodology_cell = methodology_card.rows[0].cells[0]
    _apply_cell_shading(methodology_cell, "F9FAFB")
    _set_cell_border(
        methodology_cell,
        top={"sz": 12, "color": neutral_border_hex},
        bottom={"sz": 12, "color": neutral_border_hex},
        left={"sz": 12, "color": neutral_border_hex},
        right={"sz": 12, "color": neutral_border_hex},
    )
    _set_cell_margins(methodology_cell, top=80, bottom=80, start=200, end=200)
    methodology_cell.text = ""
    methodology_title = methodology_cell.paragraphs[0]
    methodology_title.style = doc.styles["SmallCaps"]
    methodology_title.text = "Methodology"
    methodology_title.paragraph_format.space_after = Pt(6)

    def _add_methodology_line(text: str, style_name: str = "Body", space_before: float = 2, space_after: float = 2):
        paragraph = methodology_cell.add_paragraph(text, style=style_name)
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(space_after)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return paragraph

    _add_methodology_line(f"Players analyzed: {sample_size} players.")
    _add_methodology_line(
        "Percentiles and radar always use the full loaded dataset. Negative-impact metrics are auto-reversed so “higher = better”.",
    )
    _add_methodology_line("Data treatment & normalization:")
    _add_methodology_line(
        "Missing values are safely handled; per-90 and ratio metrics guard against divide-by-zero. Composite scores are built from z-scores of component stats and then normalized to a 0–100 scale. Metrics where “lower is better” (e.g., cards, conceded/xGA) are inverted so that bigger bars/areas indicate stronger performance.",
    )

    _add_spacer(6)

    _add_spacer(6)

    def _insight_context(metric_name: str, median_val) -> str:
        if median_val is None or (isinstance(median_val, float) and not math.isfinite(median_val)):
            return ""
        formatted = _format_metric_value(metric_name, median_val)
        if not formatted or formatted == "—":
            try:
                formatted = f"{float(median_val):.2f}"
            except Exception:
                return ""
        return formatted

    _add_section_heading("Visual analysis")

    analysis_intro = doc.add_paragraph(
        "Percentiles calculated on the loaded dataset (negative-impact metrics are reversed automatically).",
        style="Body",
    )
    analysis_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    analysis_intro.paragraph_format.space_before = Pt(2)
    analysis_intro.paragraph_format.space_after = Pt(6)

    chart_card = doc.add_table(rows=1, cols=1)
    chart_card.autofit = False
    chart_cell = chart_card.rows[0].cells[0]
    _set_cell_border(
        chart_cell,
        top={"sz": 14, "color": neutral_border_hex},
        bottom={"sz": 14, "color": neutral_border_hex},
        left={"sz": 14, "color": neutral_border_hex},
        right={"sz": 14, "color": neutral_border_hex},
    )
    _apply_cell_shading(chart_cell, "F9FAFB")
    _set_cell_margins(chart_cell, top=80, bottom=90, start=180, end=180)
    chart_cell.text = ""
    chart_band = chart_cell.add_paragraph(" ")
    chart_band.paragraph_format.space_after = Pt(4)
    _shade_paragraph(chart_band, accent_hex)
    chart_title = chart_cell.add_paragraph("Radar & percentile bars", style="SmallCaps")
    chart_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    chart_title.paragraph_format.space_before = Pt(0)
    chart_title.paragraph_format.space_after = Pt(4)
    chart_caption = chart_cell.add_paragraph(
        "Single-player radar alongside percentile bars across the selected metrics.",
        style="Note",
    )
    chart_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    chart_caption.paragraph_format.space_after = Pt(12)
    chart_image_para = chart_cell.add_paragraph()
    chart_image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        radar_png = make_radar_bars_png(
            df,
            player_a=player_name,
            player_b=None,
            metrics=metrics,
            color_a=accent_color,
            color_b=color_b or "#9CA3AF",
            bar_mode="percentile",
        )
        chart_width_in = min(usable_width_in - 0.25, 6.3)
        chart_width_in = max(chart_width_in, 5.5)
        chart_image_para.add_run().add_picture(radar_png, width=Inches(chart_width_in))
    except Exception:
        fallback = chart_cell.add_paragraph("Radar visualization unavailable.", style="Note")
        fallback.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    _add_section_heading("Quick insights")

    insights_table = doc.add_table(rows=1, cols=2)
    insights_table.autofit = False
    col_width_cm = max((usable_width_cm / 2) - 0.2, 6.5)
    insights_table.columns[0].width = Cm(col_width_cm)
    insights_table.columns[1].width = Cm(col_width_cm)

    standouts_cell, development_cell = insights_table.rows[0].cells
    for insight_cell in (standouts_cell, development_cell):
        _set_cell_border(
            insight_cell,
            top={"sz": 12, "color": neutral_border_hex},
            bottom={"sz": 12, "color": neutral_border_hex},
            left={"sz": 12, "color": neutral_border_hex},
            right={"sz": 12, "color": neutral_border_hex},
        )
        _set_cell_margins(insight_cell, top=90, bottom=90, start=180, end=180)

    standouts_title = standouts_cell.paragraphs[0]
    standouts_title.style = doc.styles["H2"]
    standouts_title.text = "Standouts (≥ 70th percentile)"

    for item in standouts:
        metric_label = str(item.get("metric", ""))
        descriptor = str(item.get("descriptor", ""))
        paragraph = standouts_cell.add_paragraph(style="Body")
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run("✅ ")
        paragraph.add_run(f"{metric_label}: ")
        percent_part = descriptor
        value_part = ""
        if "(" in descriptor:
            percent_part, rest = descriptor.split("(", 1)
            percent_part = percent_part.strip()
            value_part = f"({rest}".strip()
        pct_run = paragraph.add_run(percent_part.strip())
        pct_run.bold = True
        if value_part:
            paragraph.add_run(f" {value_part}")
        median_text = _insight_context(metric_label, item.get("median"))
        if median_text:
            paragraph.add_run(f" — vs dataset median {median_text}")

    development_title = development_cell.paragraphs[0]
    development_title.style = doc.styles["H2"]
    development_title.text = "Development areas (≤ 40th percentile)"

    for item in development:
        metric_label = str(item.get("metric", ""))
        descriptor = str(item.get("descriptor", ""))
        paragraph = development_cell.add_paragraph(style="Body")
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run("⚠️ ")
        paragraph.add_run(f"{metric_label}: ")
        percent_part = descriptor
        value_part = ""
        if "(" in descriptor:
            percent_part, rest = descriptor.split("(", 1)
            percent_part = percent_part.strip()
            value_part = f"({rest}".strip()
        pct_run = paragraph.add_run(percent_part.strip())
        pct_run.bold = True
        if value_part:
            paragraph.add_run(f" {value_part}")
        median_text = _insight_context(metric_label, item.get("median"))
        if median_text:
            paragraph.add_run(f" — vs dataset median {median_text}")

    _add_spacer(6)

    confidentiality_block = doc.add_table(rows=1, cols=1)
    confidentiality_block.autofit = False
    block_cell = confidentiality_block.rows[0].cells[0]
    _apply_cell_shading(block_cell, neutral_fill)
    _set_cell_border(
        block_cell,
        top={"sz": 12, "color": neutral_border_hex},
        bottom={"sz": 12, "color": neutral_border_hex},
        left={"sz": 12, "color": neutral_border_hex},
        right={"sz": 12, "color": neutral_border_hex},
    )
    _set_cell_margins(block_cell, top=120, bottom=120, start=200, end=200)
    confidentiality_para = block_cell.paragraphs[0]
    confidentiality_para.style = doc.styles["Body"]
    confidentiality_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    confidentiality_para.text = (
        "Confidential – This report contains proprietary information for authorized recipients only."
        " Do not copy, share, or distribute without written permission."
    )

    _force_update_fields(doc)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    try:
        with open(OUTPUT_FILE, "wb") as f:
            f.write(output.getbuffer())
    except Exception:
        pass

    return output
# ===================== SIDEBAR — Controls =====================
st.sidebar.header("⚙️ Settings")
up = st.sidebar.file_uploader("Upload merged Excel (WyScout + SkillCorner)", type=["xlsx"])
TOPN = st.sidebar.slider("Top N per ranking", 5, 50, 10, 1)
pos_filter = st.sidebar.text_input("Filter by Position (regex)", value="")
team_filter = st.sidebar.text_input("Filter by Team (exact match)", value="")
demo_mode = st.sidebar.checkbox("Demo mode (synthetic data)", value=False)
min_minutes = st.sidebar.number_input(
    "Minimum minutes played",
    min_value=0,
    value=0,
    step=90,
    help="Apenas os Rankings respeitarão este filtro. Radar e percentis usam o dataset completo."
)

# ===================== SIDEBAR NAVIGATION =====================

# --- Prepare df_all early so all pages can rely on it ---
df = None
if demo_mode:
    st.warning("Demo mode is ON — using a small synthetic sample for all pages.")
    df_demo = pd.DataFrame({
        "Player": ["Player A","Player B","Player C"],
        "Short Name": ["P. A","P. B","P. C"],
        "Team": ["RFS Riga","RFS Riga","RFS Riga"],
        "Position": ["FW","MF","DF"],
        "Minutes played": [900, 850, 780],
        "xG per 90":[0.35,0.20,0.05],
        "xA per 90":[0.18,0.25,0.07],
        "Successful attacking actions per 90":[3.2, 2.1, 1.4],
        "Conceded goals per 90":[0.4, 0.6, 0.8],
    })
    df = compute_composite_metrics(df_demo, DEFAULT_OFF_WEIGHTS)
elif up is not None:
    try:
        df_raw = _load_excel(up)
        df = compute_composite_metrics(df_raw, DEFAULT_OFF_WEIGHTS)
    except Exception as e:
        st.error("Erro ao carregar/calcular métricas.")
        st.exception(e)

if df is not None and not df.empty:
    df_all = _fix_npxg_block(df.copy())
page = st.sidebar.radio("📑 Pages", ["Dashboard", "Metrics Documentation", "Ferramenta de Busca"])

if page == "Metrics Documentation":
    st.title("📘 Composite Metrics Documentation")
    st.caption("Explanation of each composite metric and how it is calculated.")

    st.header("Offensive Metrics")
    st.markdown("""
**Offensive Production (per 90)**  
Weighted sum of attacking contributions such as:  
- Successful attacking actions per 90  
- xG per 90  
- xA per 90  
- Key passes per 90  
- Deep completions, progressive runs, smart passes, crosses to the goalie box, touches in box, etc.  
Weights are predefined (higher for xG/xA/key passes, lower for crosses/touches).
    """)

    st.header("Defensive Metrics")
    st.markdown("""
**Defensive Production (per 90)**  
Sum of:  
- Successful defensive actions per 90  
- Defensive duel efficiency *(duels × win%)*  
- Aerial duel efficiency *(duels × win%)*  
- Sliding tackles *(possession-adjusted)*  
- Interceptions *(possession-adjusted)*  
- Shots blocked per 90  
    """)

    st.header("Physical + Technical Composites")
    st.markdown("""
These normalize offensive and defensive production by physical output:

- **Work Rate Offensive** = *Offensive Production* ÷ *Total Distance (km per 90)*  
- **Work Rate Defensive** = *Defensive Production* ÷ *Total Distance (km per 90)*  

- **Offensive Intensity** = *Offensive Production* ÷ *(High-intensity + Running Distance, km per 90)*  
- **Defensive Intensity** = *Defensive Production* ÷ *(High-intensity + Running Distance, km per 90)*  

- **Offensive Explosion** = *Offensive Production* ÷ *Explosive Events* *(HSR, Sprints, Explosive Accelerations)*  
- **Defensive Explosion** = *Defensive Production* ÷ *Explosive Events*
    """)

    st.header("Radar Composite Metrics")
    st.markdown("""
These are normalized **[0–100]** composites built from z-scores of component stats:

- **xG Buildup**: weighted mix of xA, shot assists, npxG, key passes, deep completions, accurate passing.  
- **Creativity**: smart passes, through passes, passes to penalty area + their accuracy.  
- **Progression**: progressive passes, runs, dribbles, accelerations + accuracy of progressive actions.  
- **Defence**: defensive actions, interceptions, tackles, duels, aerials.  
- **Involvement**: combined measure of touches, passes, duels, interceptions, box presence.  
- **Discipline**: negative metric (fouls, yellows, reds).  
- **Finishing**: conversion rate, non-penalty goals, shots on target, G-xG, npxG/shot.  
- **Poaching**: balance of shot efficiency, box presence, receiving passes.  
- **Aerial Threat / Aerial Defence**: heading goals, aerial duel volume & success, interceptions, blocks.  
- **Passing Quality**: weighted mix of passing volume + accuracy across directions.  
- **Box Threat**: ratio of npxG per 90 to log(touches in box + 1).
    """)

    st.info("All metrics are scaled and normalized to ensure fair comparisons across players.")
    st.stop()

# ===================== MAIN =====================
if page == "Dashboard":
    st.title("⚽ Composite Metrics & Radar")

    st.caption("Integrated with composite metrics + physical/technical composites")



    df = None

    if demo_mode:

        st.warning("Demo mode is ON — synthetic sample loaded.")

        df_demo = pd.DataFrame({

            "Player": ["Player A","Player B","Player C"],

            "Team": ["X","Y","Z"],

            "Position": ["CF","RW","DMF"],

            "Minutes played": [900, 880, 910],

            "Successful attacking actions per 90":[5,7,2],

            "xG per 90":[0.3,0.2,0.1],

            "xA per 90":[0.2,0.4,0.05],

            "Key passes per 90":[1.2,1.5,0.6],

            "Deep completions per 90":[1.0,0.8,0.4],

            "Deep completed crosses per 90":[0.2,0.5,0.0],

            "Progressive runs per 90":[1.1,1.6,0.3],

            "Passes to penalty area per 90":[0.6,0.9,0.2],

            "Smart passes per 90":[0.4,0.5,0.2],

            "Crosses to goalie box per 90":[0.1,0.3,0.0],

            "Touches in box per 90":[3.5,4.0,1.2],

            "xG":[6.0, 4.8, 1.2],

            "Goals":[5,4,1],

            "Penalties taken":[1,0,0],

            "Shots":[20,18,9],

            "Shot assists per 90":[0.5,0.8,0.2],

            "Second assists per 90":[0.1,0.2,0.0],

            "Accurate passes %":[82,78,90],

            "Through passes per 90":[0.3,0.6,0.1],

            "Accurate smart passes, %":[55,52,60],

            "Accurate through passes, %":[42,38,50],

            "Accurate passes to penalty area, %":[40,44,35],

            "Progressive passes per 90":[3.0,3.8,1.2],

            "Accurate progressive passes, %":[68,62,70],

            "Dribbles per 90":[2.0,3.5,0.8],

            "Successful dribbles, %":[55,48,52],

            "Accelerations per 90":[1.2,1.8,0.6],

            "Successful defensive actions per 90":[4.0,2.0,6.0],

            "Defensive duels per 90":[5.0,3.0,8.0],

            "Defensive duels won, %":[60,55,68],

            "Aerial duels per 90":[2.0,1.0,3.0],

            "Aerial duels won, %":[50,45,62],

            "PAdj Sliding tackles":[0.3,0.1,0.7],

            "PAdj Interceptions":[0.8,0.4,1.2],

            "Shots blocked per 90":[0.2,0.1,0.5],

            "Passes per 90":[35,42,55],

            "Received passes per 90":[12,14,10],

            "Touches per 90":[50,48,62],

            "Passes to final third per 90":[2.2,1.9,1.5],

            "Accurate passes to final third, %":[70,64,72],

            "Forward passes per 90":[12,14,10],

            "Accurate forward passes, %":[78,75,82],

            "Long passes per 90":[3.0,2.0,5.0],

            "Accurate long passes, %":[55,48,62],

            "Lateral passes per 90":[7.0,8.0,6.0],

            "Accurate lateral passes, %":[90,88,92],

            "Back passes per 90":[4.0,6.0,5.0],

            "Accurate back passes, %":[94,96,95],

            "Head goals per 90":[0.05,0.02,0.03],

            "Goal conversion, %":[18,15,8],

            "Shots on target, %":[45,42,38],

            "Fouls per 90":[1.5,1.8,2.2],

            "Yellow cards per 90":[0.2,0.15,0.25],

            "Red cards per 90":[0.01,0.0,0.02],

            "Distance P90":[10000,9800,10500],

            "Running Distance P90":[8000,7800,8200],

            "HI Distance P90":[1500,1300,1100],

            "HSR Distance P90":[600,550,500],

            "Sprint Distance P90":[250,220,200],

            "HSR Count P90":[40,35,30],

            "Sprint Count P90":[15,12,10],

            "Explosive Acceleration to HSR Count P90":[3,2,2],

            "Explosive Acceleration to Sprint Count P90":[1,1,1],

        })

        df = compute_composite_metrics(df_demo, DEFAULT_OFF_WEIGHTS)
if page == "Ferramenta de Busca":
    # --- Garantia: só continue se df_all existir e estiver pronto ---
    try:
        _df_all_ready = hasattr(df_all, "columns") and len(df_all.columns) > 0
    except NameError:
        _df_all_ready = False
    if not _df_all_ready:
        st.info("Envie o Excel combinado na barra lateral (ou ative o modo Demo) para usar a Ferramenta de Busca.")
        st.stop()

    st.title("🔎 Ferramenta de Busca")
    st.caption("Filtre jogadores por **percentis mínimos** em um *position preset*. Os percentis são calculados no **dataset completo** (com inversão de métricas negativas), enquanto o filtro de **minutos** é aplicado ao final.")

    # Seleção de preset e configuração dos limiares
    preset_name = st.selectbox("Position preset", options=list(PRESETS.keys()))
    metrics_for_preset = PRESETS[preset_name]

    use_global = st.checkbox("Usar um único percentil mínimo para todas as métricas", value=True)
    global_min = st.slider("Percentil mínimo (todas as métricas)", 0, 100, 70, help="Aplica-se somente se a opção acima estiver marcada.")

    # Limiares por métrica
    thresholds = {}
    if use_global:
        thresholds = {m: global_min for m in metrics_for_preset}
    else:
        st.markdown("#### Limiares por métrica")
        for m in metrics_for_preset:
            thresholds[m] = st.slider(f"{m}", 0, 100, 70)

        # Cálculo de percentis por métrica (dataset completo) — rank-based (0–100)
    def _percentile_rank_series(s: pd.Series, ascending: bool) -> pd.Series:
        s_num = pd.to_numeric(s, errors="coerce")
        mask = s_num.notna()
        n = int(mask.sum())
        out = pd.Series(np.nan, index=s.index)
        if n <= 1:
            return out
        r = s_num[mask].rank(ascending=ascending, method="average")
        out.loc[mask] = 100.0 * (n - r) / (n - 1)
        return out

    percent_cols = {}
    for m in metrics_for_preset:
        if m in df_all.columns:
            ascending = (m in NEGATE_METRICS)  # métricas negativas: menor é melhor
            percent_cols[m] = _percentile_rank_series(df_all[m], ascending=ascending)

    res = df_all.copy()
    for m, col in percent_cols.items():
        res[f"{m} (pct)"] = col

    # Colunas base e lista de colunas de percentil
    base_cols = [c for c in ["Player", "Short Name", "Team", "Position", "Minutes played", "Minutes"] if c in res.columns]
    pct_col_names = [f"{m} (pct)" for m in metrics_for_preset if f"{m} (pct)" in res.columns]


    # Aplica filtros de minutes, team e posição já existentes
    if "Minutes played" in res.columns:
        min_col = "Minutes played"
    elif "Minutes" in res.columns:
        min_col = "Minutes"
    else:
        min_col = None

    if min_col is not None and isinstance(min_minutes, (int, float)):
        res = res[pd.to_numeric(res[min_col], errors="coerce").fillna(0) >= float(min_minutes)]

    if team_filter:
        if "Team" in res.columns:
            res = res[res["Team"].astype(str) == team_filter]

    if pos_filter:
        if "Position" in res.columns:
            res = res[res["Position"].astype(str).str.contains(pos_filter)]

    # Filtro por percentis mínimos (todas as métricas do preset)
    mask = pd.Series(True, index=res.index)
    for m in metrics_for_preset:
        pct_col = f"{m} (pct)"
        if pct_col in res.columns:
            mask &= (pd.to_numeric(res[pct_col], errors="coerce") >= thresholds.get(m, 0))
        else:
            # Se a métrica não existir, o jogador falha no critério
            mask &= False
    res = res[mask]

    # Ordenação por média dos percentis do preset (opcional)
    pct_cols = [f"{m} (pct)" for m in metrics_for_preset if f"{m} (pct)" in res.columns]
    if pct_cols:
        res["Media pct (preset)"] = res[pct_cols].mean(axis=1)
        res = res.sort_values("Media pct (preset)", ascending=False)

    # Mostra resultados
    show_cols = base_cols + pct_col_names
    show_cols = [c for c in show_cols if c in res.columns]
    st.markdown(f"**Jogadores encontrados:** {len(res)}")
    if len(show_cols) == 0:
        st.info("Nenhuma coluna válida para exibir.")
    else:
        st.dataframe(res[show_cols].reset_index(drop=True))

    # Download
    csv_bytes = res[show_cols].to_csv(index=False).encode("utf-8") if show_cols else b""
    st.download_button("Baixar resultados (CSV)", data=csv_bytes, file_name="busca_percentis.csv", mime="text/csv")

else:
    if up is None:
        st.info("Upload your merged Excel on the left panel to begin (or enable Demo mode).")
    else:
        with st.spinner("Loading data and computing metrics…"):
            try:
                df_raw = _load_excel(up)
                df = compute_composite_metrics(df_raw, DEFAULT_OFF_WEIGHTS)
            except Exception as e:
                st.error("There was an error while computing metrics.")
                st.exception(e)

if df is None or df.empty:
    st.stop()

# Enforce presence of Minutes played
if "Minutes played" not in df.columns:
    st.error("Arquivo não possui a coluna 'Minutes played'. Renomeie a coluna para exatamente 'Minutes played' e reenvié.")
    st.stop()

# Keep global (df_all) for all calculations; df_view only controls Rankings listing
df_all = df.copy()
df_all = _fix_npxg_block(df_all)

df_view = df_all[df_all["Minutes played"].fillna(0) >= int(min_minutes)].copy()
st.caption(
    f"Filtro de minutos afeta apenas os **Rankings** (mostra {df_view.shape[0]} de {df_all.shape[0]} jogadores). "
    "Cálculos de radar e percentis permanecem sobre o dataset completo."
)

# ===================== KPIs =====================
st.subheader("Overview")
col1, col2, col3, col4 = st.columns(4)
with col1: st.markdown(f"<div class='metric-card'><div class='subtle small'>Players</div><h3>{df_all.shape[0]}</h3></div>", unsafe_allow_html=True)
with col2: st.markdown(f"<div class='metric-card'><div class='subtle small'>Teams</div><h3>{df_all['Team'].nunique() if 'Team' in df_all.columns else '—'}</h3></div>", unsafe_allow_html=True)
with col3: st.markdown(f"<div class='metric-card'><div class='subtle small'>Positions</div><h3>{df_all['Position'].nunique() if 'Position' in df_all.columns else '—'}</h3></div>", unsafe_allow_html=True)
with col4: st.markdown(f"<div class='metric-card'><div class='subtle small'>Columns</div><h3>{df_all.shape[1]}</h3></div>", unsafe_allow_html=True)

# ===================== Rankings =====================
st.markdown("<div class='section'>Rankings</div>", unsafe_allow_html=True)

def _leaderboard(metric):
    d = df_view.copy()
    if d.empty:
        st.info("Nenhum jogador atende ao filtro de minutos para exibição no ranking.")
        return
    if team_filter:
        d = d[d["Team"].astype(str) == team_filter]
    if pos_filter:
        d = d[d["Position"].astype(str).str.contains(pos_filter)]
    if metric not in d.columns:
        st.warning(f"Métrica {metric} não encontrada no dataset.")
        return
    id_cols = [c for c in ID_COLS_CANDIDATES if c in d.columns]
    cols = id_cols + [
        "Work Rate Offensive","Offensive Intensity","Offensive Explosion",
        "Work Rate Defensive","Defensive Intensity","Defensive Explosion",
        "Creativity","Progression","Defence","Passing Quality","Aerial Defence",
        "Involvement","Discipline","xG Buildup","Box Threat","Finishing",
        "Poaching","Aerial Threat","npxG per 90","npxG per Shot","G-xG",
    ]
    present = [c for c in cols if c in d.columns]
    out = d.dropna(subset=[metric]).sort_values(metric, ascending=False).head(TOPN)
    st.dataframe(out[present], use_container_width=True)

t1, t2, t3, t4, t5, t6 = st.tabs([
    "Work Rate Offensive","Offensive Intensity","Offensive Explosion",
    "Work Rate Defensive","Defensive Intensity","Defensive Explosion",
])
with t1: _leaderboard("Work Rate Offensive")
with t2: _leaderboard("Offensive Intensity")
with t3: _leaderboard("Offensive Explosion")
with t4: _leaderboard("Work Rate Defensive")
with t5: _leaderboard("Defensive Intensity")
with t6: _leaderboard("Defensive Explosion")

# ===================== Radar Generator + Ranking Bars =====================
st.markdown("<div class='section'>Radar Generator</div>", unsafe_allow_html=True)

def _merge_presets(preset_names: list[str], df: pd.DataFrame) -> list[str]:
    merged = []
    for p in preset_names:
        for m in PRESETS.get(p, []):
            if m in df.columns and df[m].notna().any() and m not in merged:
                merged.append(m)
    return merged

colA, colB = st.columns([1, 2])
with colA:
    all_presets = list(PRESETS.keys())
    selected_presets = st.multiselect(
        "Position presets (choose up to 3)",
        options=all_presets,
        default=[all_presets[0]]
    )
    if len(selected_presets) > 3:
        st.warning("You selected more than 3 presets; only the first 3 will be used.")
        selected_presets = selected_presets[:3]

    metrics_from_presets = _merge_presets(selected_presets, df_all)

    metrics_all = sorted([
        c for c in df_all.columns
        if c not in ID_COLS_CANDIDATES and pd.api.types.is_numeric_dtype(df_all[c])
    ])

    default_metrics = metrics_from_presets[:16] if metrics_from_presets else []
    metrics_sel = st.multiselect(
        "Metrics in radar (max 16)",
        options=metrics_all,
        default=default_metrics,
        help="You can add/remove metrics. If multiple presets are selected, duplicates are removed automatically."
    )
    if len(metrics_sel) > 16:
        st.info(f"You selected {len(metrics_sel)} metrics; only the first 16 will be plotted.")

    players = sorted(df_all["Player"].dropna().unique().tolist()) if "Player" in df_all.columns else []
    p1 = st.selectbox("Player A", players)
    p2 = st.selectbox("Player B (optional)", ["—"] + players)
    color_a = st.color_picker("Color A", "#2A9D8F")
    color_b = st.color_picker("Color B", "#E76F51")
# Download button for combined PNG
if p1 and metrics_sel:
    png_buf = make_radar_bars_png(
        df_all,
        p1,
        None if p2 == "—" else p2,
        metrics_sel,
        color_a,
        color_b,
    )
    st.download_button(
        "⬇️ Download Radar + Barras (PNG)",
        data=png_buf.getvalue(),
        file_name="radar_barras.png",
        mime="image/png",
    )
with colB:
    if p1 and metrics_sel:
        plot_radar(df_all, p1, None if p2 == "—" else p2, metrics_sel, color_a, color_b)
        render_metric_rank_bars(df_all, p1, metrics_sel, None if p2 == "—" else p2)

# ===================== Export =====================
st.markdown("<div class='section'>Export</div>", unsafe_allow_html=True)
st.download_button(
    "Download full CSV",
    df_all.to_csv(index=False).encode("utf-8"),
    file_name="composite_metrics_base.csv",
    mime="text/csv",
)

st.markdown("#### Relatório DOCX (Radar + Percentis)")

if "docx_competition" not in st.session_state:
    st.session_state.docx_competition = ""
if "docx_season" not in st.session_state:
    st.session_state.docx_season = ""
if "docx_show_competition" not in st.session_state:
    st.session_state.docx_show_competition = False
if "docx_show_season" not in st.session_state:
    st.session_state.docx_show_season = False

prompt_col_a, prompt_col_b = st.columns(2)
with prompt_col_a:
    if st.button("Set competition", key="docx_set_competition"):
        st.session_state.docx_show_competition = True
    if st.session_state.get("docx_show_competition", False):
        st.text_input(
            "Competition",
            key="docx_competition",
            help="Use this field to override the competition shown in the report.",
        )
with prompt_col_b:
    if st.button("Set season", key="docx_set_season"):
        st.session_state.docx_show_season = True
    if st.session_state.get("docx_show_season", False):
        st.text_input(
            "Season / Year",
            key="docx_season",
            help="Provide the season or year label that appears as the report date.",
        )

col_doc_a, col_doc_b = st.columns(2)
with col_doc_a:
    player_photo_upload = st.file_uploader(
        "Foto do atleta (opcional)",
        type=["png", "jpg", "jpeg"],
        key="docx_player_photo",
        help="Imagem será posicionada no cabeçalho do relatório.",
    )
with col_doc_b:
    team_logo_upload = st.file_uploader(
        "Escudo do time (opcional)",
        type=["png", "jpg", "jpeg"],
        key="docx_team_logo",
    )

valid_metrics_for_docx = [m for m in metrics_sel if m in df_all.columns] if p1 else []
if p1 and len(valid_metrics_for_docx) >= 3:
    try:
        report_buf = build_player_report_docx(
            df_all,
            player_name=p1,
            metrics=valid_metrics_for_docx,
            color_a=color_a,
            color_b=color_b,
            competition_level_override=st.session_state.get("docx_competition"),
            player_photo=player_photo_upload.getvalue() if player_photo_upload else None,
            team_logo=team_logo_upload.getvalue() if team_logo_upload else None,
            report_date=st.session_state.get("docx_season"),
        )
        st.download_button(
            "⬇️ Baixar relatório DOCX",
            data=report_buf.getvalue(),
            file_name=f"{_slugify_filename(p1)}_relatorio_radar.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error("Não foi possível gerar o relatório DOCX.")
        st.exception(e)
elif p1 and metrics_sel:
    st.info("Selecione pelo menos 3 métricas válidas para gerar o relatório DOCX.")
else:
    st.info("Escolha um jogador e as métricas desejadas para habilitar o relatório DOCX.")


# --- Revised presets (auto-generated) ---
PRESETS = {'aerial_duels': ['Aerial Threat',
                  'Aerial Defence',
                  'Aerial duels per 90',
                  'Aerial duels won, %',
                  'Head goals per 90',
                  'Involvement',
                  'Defence'],
 'attacking_midfielder': ['Creativity',
                          'Progression',
                          'xA per 90',
                          'Key passes per 90',
                          'Smart passes per 90',
                          'Deep completions per 90',
                          'Work Rate Offensive',
                          'Offensive Intensity',
                          'Offensive Explosion'],
 'center_back': ['Defence',
                 'Aerial Defence',
                 'Aerial duels per 90',
                 'Aerial duels won, %',
                 'Defensive duels per 90',
                 'Defensive duels won, %',
                 'PAdj Interceptions',
                 'PAdj Sliding tackles',
                 'Passes to final third per 90',
                 'Accurate passes %',
                 'Work Rate Defensive',
                 'Defensive Intensity',
                 'Defensive Explosion'],
 'central_midfielder': ['Progression',
                        'Passing Quality',
                        'PAdj Interceptions',
                        'Successful defensive actions per 90',
                        'Work Rate Defensive',
                        'Defensive Intensity',
                        'Passes to final third per 90',
                        'Progressive runs per 90',
                        'Deep completions per 90'],
 'counter_attack': ['Progression',
                    'Accelerations per 90',
                    'Dribbles per 90',
                    'Successful dribbles, %',
                    'npxG per 90',
                    'Finishing',
                    'Box Threat'],
 'crossing': ['Crossing',
              'Accurate crosses, %',
              'Deep completed crosses per 90',
              'xA per 90',
              'Shot assists per 90',
              'Creativity',
              'Passing Quality'],
 'defensive_actions': ['Defence',
                       'Aerial Defence',
                       'PAdj Interceptions',
                       'Successful defensive actions per 90',
                       'Defensive duels won, %',
                       'Shots blocked per 90',
                       'Discipline'],
 'defensive_midfielder': ['Defence',
                          'Progression',
                          'Discipline',
                          'PAdj Interceptions',
                          'Successful defensive actions per 90',
                          'Work Rate Defensive',
                          'Defensive Intensity',
                          'Defensive Explosion',
                          'Defensive duels won, %',
                          'Aerial duels won, %'],
 'forward': ['npxG per 90',
             'xG per 90',
             'Shots per 90',
             'Shots on target, %',
             'xA per 90',
             'Key passes per 90',
             'Touches in box per 90',
             'Progressive runs per 90',
             'Deep completions per 90',
             'Finishing',
             'Poaching',
             'Aerial Threat',
             'Work Rate Offensive',
             'Offensive Intensity',
             'Offensive Explosion'],
 'full_back': ['Progression',
               'Creativity',
               'Passing Quality',
               'Aerial Defence',
               'Deep completed crosses per 90',
               'Progressive runs per 90',
               'Crosses per 90',
               'Work Rate Defensive',
               'Defensive Intensity',
               'Defensive Explosion'],
 'general_summary': ['Involvement', 'Creativity', 'Box Threat', 'Discipline'],
 'playmaking_build_up': ['Successful attacking actions per 90',
                         'Deep completions per 90',
                         'Key passes per 90',
                         'Discipline',
                         'Creativity',
                         'Passing Quality',
                         'Progression'],
 'shooting': ['npxG per 90',
              'npxG per Shot',
              'Finishing',
              'Goal conversion, %',
              'Shots on target, %',
              'G-xG',
              'Box Threat'],
 'striker': ['Involvement',
             'npxG per 90',
             'npxG per Shot',
             'Finishing',
             'Poaching',
             'Aerial Threat',
             'Box Threat',
             'Touches in box per 90',
             'Shots per 90',
             'Shots on target, %'],
 'winger': ['Creativity',
            'Progression',
            'Dribbles per 90',
            'Dribbles won, %',
            'Crosses per 90',
            'Accurate crosses, %',
            'Deep completed crosses per 90',
            'Progressive runs per 90',
            'xA per 90',
            'Key passes per 90',
            'Work Rate Offensive',
            'Offensive Intensity',
            'Offensive Explosion',
            'Successful dribbles, %']}


# --- Guarantee fusion of playmaking + build_up into playmaking_build_up (runs after any PRESETS re-definitions) ---
def _radar_norm_key__pm_bu(s: str) -> str:
    return (
        str(s).strip().lower().replace("-", "_").replace(" ", "_")
    )

def _radar_dedup_keep_order__pm_bu(seq):
    seen = set(); out = []
    for x in seq:
        if x not in seen:
            out.append(x); seen.add(x)
    return out

try:
    _keys_map = { _radar_norm_key__pm_bu(k): k for k in list(PRESETS.keys()) }
    _k_play  = _keys_map.get("playmaking")
    _k_build = _keys_map.get("build_up") or _keys_map.get("buildup")
    if _k_play and _k_build:
        _merged = _radar_dedup_keep_order__pm_bu(PRESETS[_k_play] + PRESETS[_k_build])
        _CORE = {"Progression", "Creativity", "Passing Quality", "xG Buildup", "Defence", "Involvement"}
        _core_cnt = 0; _merged_limited = []
        for m in _merged:
            if m in _CORE:
                if _core_cnt >= 3:
                    continue
                _core_cnt += 1
            _merged_limited.append(m)
        PRESETS["playmaking_build_up"] = _merged_limited[:10]
        del PRESETS[_k_play]
        del PRESETS[_k_build]
except Exception as _e:
    # Do not break the app if fusion fails for any reason
    pass


# --- A4 PDF export: horizontal percentiles-by-cohort bars with P50/P80 guides ---
import numpy as np, pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

def _percentile_rank_by_cohort_safe(df: pd.DataFrame, metric: str,
                                    league_col: str = "League", pos_col: str = "Position",
                                    season_col: str | None = None) -> pd.Series:
    s = pd.to_numeric(df.get(metric, pd.Series(index=df.index, dtype=float)), errors="coerce")
    mask = s.notna()
    out = pd.Series(np.nan, index=df.index)

    group_cols = [c for c in [league_col, pos_col, season_col] if c and c in df.columns]
    if not group_cols:
        def _pct_rank(x):
            n = x.shape[0]
            if n <= 1: return pd.Series(np.nan, index=x.index)
            r = x.rank(ascending=False, method="average")
            return 100.0 * (n - r) / (n - 1)
        out.loc[mask] = _pct_rank(s[mask])
        return out

    grp_keys = df.loc[mask, group_cols].apply(tuple, axis=1)
def _choose_default_blocks(df: pd.DataFrame) -> dict:
    blocks = {}
    cand_off = [c for c in [
        "npxG per 90","Non-penalty goals per 90","Shots per 90","Shots on target, %",
        "xA per 90","Key passes per 90","Smart passes per 90","Deep completions per 90",
        "Touches in box per 90","Box Threat"
    ] if c in df.columns]
    if cand_off: blocks["Perfil ofensivo"] = cand_off[:8]

    cand_prog = [c for c in [
        "Passes to final third per 90","Passes to penalty area per 90","Progressive runs per 90",
        "Progressive passes per 90","Accurate passes, %"
    ] if c in df.columns]
    if cand_prog: blocks["Progressão"] = cand_prog[:8]

    cand_def = [c for c in [
        "Successful defensive actions per 90","PAdj Interceptions","Defensive duels won, %",
        "Aerial duels won, %","Recoveries per 90","Discipline"
    ] if c in df.columns]
    if cand_def: blocks["Defesa/Pressão"] = cand_def[:8]

    return blocks

def export_player_pdf_a4_bars(df: pd.DataFrame, player_name: str,
                              cohort_filters: dict | None = None,
                              blocks: dict | None = None,
                              output_path: str = "player_report.pdf",
                              league_col: str = "League", pos_col: str = "Position",
                              season_col: str | None = None, minutes_col: str = "Minutes played"):
    """
    Creates a 1–2 page A4 landscape PDF with horizontal percentile bars by cohort and cut lines at P50/P80.
    - df: your master dataframe
    - player_name: exact name from df['Player']
    - cohort_filters: e.g., {"League": "Championship", "Position": "CB"} to restrict coorte
    - blocks: dict of {"Section title": [metric1, metric2, ...]}. If None, a sensible default is chosen.
    """
    if "Player" not in df.columns:
        raise ValueError("DataFrame must have a 'Player' column.")

    d = df.copy()
    if cohort_filters:
        for k, v in cohort_filters.items():
            if k in d.columns:
                d = d[d[k] == v]

    cohort_df = d if len(d) >= 3 else df

    if blocks is None:
        blocks = _choose_default_blocks(cohort_df)

    if player_name in d["Player"].values:
        row = d[d["Player"] == player_name].iloc[0]
    else:
        row = df[df["Player"] == player_name].iloc[0]
        d = df

    try:
        minutes = float(row.get(minutes_col, np.nan))
    except Exception:
        minutes = np.nan

    pos = row.get(pos_col) if pos_col in row.index else None
    league = row.get(league_col) if league_col in row.index else None

    plot_blocks = []
    for title, metrics in blocks.items():
        avail = [m for m in metrics if m in cohort_df.columns]
        if not avail: 
            continue
        pct_map = {}
        raw_map = {}
        for m in avail:
            pct_series = _percentile_rank_by_cohort_safe(cohort_df, m, league_col, pos_col, season_col)
            pct_series = pct_series.reindex(cohort_df.index)
            val_pct = np.nan
            if row.name in pct_series.index and pd.notna(pct_series.loc[row.name]):
                val_pct = float(pct_series.loc[row.name])
            pct_map[m] = val_pct
            try:
                raw_map[m] = float(pd.to_numeric(d.loc[row.name][m], errors="coerce")) if (row.name in d.index and m in d.columns) else np.nan
            except Exception:
                raw_map[m] = np.nan
        kept = [m for m in avail if not np.isnan(pct_map[m])]
        if kept:
            plot_blocks.append((title, kept, pct_map, raw_map))

    if not plot_blocks:
        raise ValueError("No metrics available to plot for the selected player/cohort.")

    a4_landscape = (11.69, 8.27)
    with PdfPages(output_path) as pdf:
        per_page = 3
        for page_idx in range(0, len(plot_blocks), per_page):
            page_blocks = plot_blocks[page_idx:page_idx+per_page]

            fig = plt.figure(figsize=a4_landscape)
            fig.suptitle(
                f"{player_name} — {pos or ''} — {league or ''}  |  Percentis por coorte"
                + (f"  |  Minutos: {int(minutes)}" if not np.isnan(minutes) else ""),
                fontsize=14, y=0.98
            )

            top = 0.90; left = 0.08; right = 0.95; vspace = 0.26
            for bi, (title, mets, pmap, rmap) in enumerate(page_blocks):
                ax = fig.add_axes([left, top - (bi+1)*vspace + 0.03, right-left, vspace-0.06])
                vals = [pmap[m] for m in mets]
                y = np.arange(len(mets))
                ax.barh(y, vals)
                ax.set_yticks(y, labels=mets, fontsize=9)
                ax.set_xlim(0, 100)
                ax.set_xlabel("Percentil (0–100)")
                ax.set_title(title, loc="left", fontsize=12)

                ax.axvline(50, linestyle="--", linewidth=1)
                ax.axvline(80, linestyle="--", linewidth=1)

                for yi, m in enumerate(mets):
                    rv = rmap[m]
                    if not (rv is None or np.isnan(rv)):
                        ax.text(vals[yi] + 1, yi, f"{rv:.2f}", va="center", fontsize=8)

                ax.grid(True, axis="x", linewidth=0.3, alpha=0.4)
            pdf.savefig(fig, bbox_inches="tight")
            plt.close(fig)
